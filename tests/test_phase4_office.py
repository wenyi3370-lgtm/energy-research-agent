from __future__ import annotations

import json
import tempfile
import unittest
import zipfile
from pathlib import Path

from enterprise_energy_research.artifacts.excel import ExcelMasterFrozenPublisher
from enterprise_energy_research.artifacts.word import FrozenWordPublisher
from enterprise_energy_research.domain.enums import ArtifactType, RunStatus
from enterprise_energy_research.domain.ids import new_sortable_id
from enterprise_energy_research.domain.models import ExtractedEvidenceBatch, RunManifest
from enterprise_energy_research.evidence.freeze import FreezeService
from enterprise_energy_research.evidence.store import EvidenceStore
from enterprise_energy_research.graph.phase3_runner import Phase3Runner
from enterprise_energy_research.graph.state import ResearchState
from enterprise_energy_research.settings import load_yaml


ROOT = Path(__file__).resolve().parents[1]


class Phase4OfficeTests(unittest.TestCase):
    def _bundle(self, temp: str):
        raw = json.loads((ROOT / "tests" / "fixtures" / "normal_manufacturer.json").read_text(encoding="utf-8"))
        company = raw[0]["entities"][0]["canonical_name"]
        run_id, request_id = new_sortable_id("RUN"), new_sortable_id("REQ")
        store = EvidenceStore(Path(temp) / "evidence.sqlite3")
        store.create_run(RunManifest(run_id=run_id, request_id=request_id, status=RunStatus.RUNNING, config_hash="fixture", code_version="0.9.1", model_gateway={"mode": "fixture"}))
        state, manifest, _ = Phase3Runner(store, load_yaml(ROOT / "config" / "enterprise_rules.yaml")).process_batches(
            ResearchState(run_id=run_id, request_id=request_id, status=RunStatus.RUNNING), company,
            [ExtractedEvidenceBatch.model_validate(item) for item in raw], output_dir=Path(temp) / "freeze",
        )
        return FreezeService(store).load_bundle(state.freeze_id), manifest

    def test_excel_master_publisher_creates_expected_sheets(self) -> None:
        from openpyxl import load_workbook
        with tempfile.TemporaryDirectory() as temp:
            bundle, manifest = self._bundle(temp)
            binding = next(item for item in manifest.artifacts if item.type == ArtifactType.EXCEL)
            target = Path(temp) / "report.xlsx"
            result = ExcelMasterFrozenPublisher().publish(bundle, binding, target)
            self.assertEqual(result.status, "published")
            workbook = load_workbook(target, read_only=False, data_only=False)
            self.assertEqual(workbook.sheetnames, ["运行清单", "企业实体", "生产基地", "产品", "证据主表", "来源", "图片证据", "数据缺口", "能源画像", "合作机会"])
            self.assertFalse(workbook["证据主表"].sheet_view.showGridLines)
            self.assertIsNotNone(workbook["证据主表"].freeze_panes)

    def test_word_publisher_contains_real_toc_and_page_fields(self) -> None:
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        with tempfile.TemporaryDirectory() as temp:
            bundle, manifest = self._bundle(temp)
            binding = next(item for item in manifest.artifacts if item.type == ArtifactType.WORD)
            target = Path(temp) / "report.docx"
            result = FrozenWordPublisher().publish(bundle, binding, target)
            self.assertEqual(result.status, "published")
            with zipfile.ZipFile(target) as archive:
                document_xml = archive.read("word/document.xml").decode("utf-8")
                footer_xml = archive.read("word/footer1.xml").decode("utf-8")
            self.assertIn("TOC", document_xml)
            self.assertIn("PAGE", footer_xml)
            self.assertIn("tblLayout", document_xml)
            self.assertNotIn('w:val="TableGrid"', document_xml)
            # consulting footer: source + date + bias note (footer1.xml)
            self.assertIn("数据来源：公开渠道已核验证据", footer_xml)
            self.assertIn("偏差说明", footer_xml)
            # figure source notes live next to figures in the body
            self.assertIn("数据来源：", document_xml)
            report = Document(target)
            self.assertTrue(report.tables)
            for table in report.tables:
                self.assertEqual(table.alignment, WD_TABLE_ALIGNMENT.CENTER)
                for row in table.rows:
                    for cell in row.cells:
                        self.assertEqual(cell.vertical_alignment, WD_CELL_VERTICAL_ALIGNMENT.CENTER)
                        for paragraph in cell.paragraphs:
                            self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
                            self.assertEqual(paragraph.paragraph_format.first_line_indent, 0)
                            self.assertEqual(paragraph.paragraph_format.left_indent, 0)
                            self.assertEqual(paragraph.paragraph_format.right_indent, 0)
            asset_root = target.parent / "report_assets"
            visual_manifest = asset_root / "visual_manifest.json"
            self.assertTrue(visual_manifest.is_file())
            payload = json.loads(visual_manifest.read_text(encoding="utf-8"))
            self.assertEqual(payload["visual_system"], "diagram-design")
            self.assertEqual(payload["theme"], "enterprise-consulting-diagram-design")
            self.assertEqual(payload["schema_version"], "2.0")
            figures = asset_root / "figures"
            png_visuals = [
                item for item in payload["visuals"]
                if (figures / f"{item['visual_id']}.png").is_file()
            ]
            # every visual with a rendered PNG is embedded in the document
            self.assertGreaterEqual(document_xml.count("<w:drawing>"), len(png_visuals))
            for item in payload["visuals"]:
                self.assertIn(item["visual_type"], {
                    "line", "bar", "radar", "quadrant", "scatter", "treemap", "timeline",
                    "process", "data_flow", "sankey", "gantt", "pyramid", "tree",
                    "fishbone", "architecture", "journey", "kpi_cards", "table",
                })
                self.assertTrue(item["decision_question"])
                self.assertTrue(item["business_thesis"])
                self.assertTrue((figures / f"{item['visual_id']}.html").is_file())
                self.assertTrue((figures / f"{item['visual_id']}.svg").is_file())
            # QA report exists and is separate from the user document
            qa_path = asset_root / "publication_qa_report.json"
            self.assertTrue(qa_path.is_file())
            qa = json.loads(qa_path.read_text(encoding="utf-8"))
            self.assertEqual(qa["freeze_id"], bundle.freeze.freeze_id)
            self.assertNotIn("QA", document_xml)
            # narrative artifact drives the document
            self.assertTrue((asset_root / "narrative.json").is_file())


if __name__ == "__main__":
    unittest.main()
