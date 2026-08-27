"""T2 门禁修复回归：渠道回填（图4/章八）与表引用补全（表1-1/verify[9]）。"""
from __future__ import annotations

import csv
from pathlib import Path

from energy_research_agent.agent.market_production import MarketProductionPipeline


def _write_csv(path: Path, header: list[str], rows: list[list[str]]) -> None:
    with path.open("w", encoding="utf-8", newline="") as fh:
        writer = csv.writer(fh)
        writer.writerow(header)
        writer.writerows(rows)


def _read(path: Path) -> list[dict[str, str]]:
    with path.open(encoding="utf-8-sig", newline="") as fh:
        return list(csv.DictReader(fh))


# ---- 渠道回填：05.channel 全空时用 06 登记值按型号/品牌回填 ----

def test_backfill_pricing_channel_from_channel_service(tmp_path: Path) -> None:
    _write_csv(
        tmp_path / "05_Pricing_Channel.csv",
        ["pricing_id", "brand", "exact_model", "channel", "channel_type"],
        [
            ["PR-001", "隆基", "Hi-MO X10", "", ""],
            ["PR-002", "优能", "", "", ""],
            ["PR-003", "华为", "LUNA2000", "官网直销", "直销"],  # 已填不回填
        ],
    )
    _write_csv(
        tmp_path / "06_Channel_Service.csv",
        ["brand", "exact_model", "online_channel", "offline_channel", "installation_service"],
        [
            ["隆基", "Hi-MO X10", "品牌官网", "分销商", "EPC"],
            ["优能", "", "", "安装商", ""],
            ["华为", "LUNA2000", "官网直销", "", ""],
        ],
    )
    pipeline = MarketProductionPipeline(tmp_path, None, scripts_dir=tmp_path)
    pipeline._backfill_pricing_channels()

    rows = _read(tmp_path / "05_Pricing_Channel.csv")
    by_model = {r["exact_model"] or r["brand"]: r["channel"] for r in rows}
    assert by_model["Hi-MO X10"] == "品牌官网、分销商、EPC"  # 型号级匹配
    assert by_model["优能"] == "安装商"  # 品牌级匹配
    assert by_model["LUNA2000"] == "官网直销"  # 已有值不被覆盖


def test_backfill_no_channel_column_is_noop(tmp_path: Path) -> None:
    _write_csv(
        tmp_path / "05_Pricing_Channel.csv",
        ["pricing_id", "brand"],  # 无 channel 列
        [["PR-001", "隆基"]],
    )
    _write_csv(
        tmp_path / "06_Channel_Service.csv",
        ["brand", "online_channel"],
        [["隆基", "品牌官网"]],
    )
    pipeline = MarketProductionPipeline(tmp_path, None, scripts_dir=tmp_path)
    pipeline._backfill_pricing_channels()  # 不抛异常
    assert _read(tmp_path / "05_Pricing_Channel.csv")[0]["brand"] == "隆基"


# ---- 表引用补全：缺失的（见表1-1）追加到本章正文段尾 ----

def _make_report_with_missing_table_ref(tmp_path: Path) -> Path:
    from docx import Document

    doc = Document()
    doc.add_heading("一、执行摘要", level=1)
    doc.add_paragraph("本章概述德国户用储能市场的核心结论与进入建议，供决策参考。")
    doc.add_paragraph("表1-1 执行摘要关键数据与来源")  # 表题，但正文无（见表1-1）
    doc.add_heading("二、调研范围与方法", level=1)
    doc.add_paragraph("本章说明采集边界与证据体系，为后续章节提供可追溯底座。")
    doc.add_paragraph("表2-1 采集任务登记口径")
    doc.add_paragraph("本章关键数据（见表2-1）。")  # 表2-1 已有引用
    report = tmp_path / "report.docx"
    doc.save(report)
    return report


def test_ensure_table_references_adds_missing_ref(tmp_path: Path) -> None:
    report = _make_report_with_missing_table_ref(tmp_path)
    pipeline = MarketProductionPipeline(tmp_path, None, scripts_dir=tmp_path)
    pipeline._ensure_table_references(report)

    from docx import Document

    doc = Document(report)
    body = "".join(p.text for p in doc.paragraphs)
    assert "（见表1-1）" in body  # 缺失的被补上
    assert body.count("（见表2-1）") == 1  # 已有的不重复追加
    # 引用落在执行摘要章的正文段（而非表题/标题段）
    first_chapter_para = doc.paragraphs[1].text
    assert "（见表1-1）" in first_chapter_para


def test_ensure_table_references_noop_when_all_present(tmp_path: Path) -> None:
    from docx import Document

    doc = Document()
    doc.add_heading("一、执行摘要", level=1)
    doc.add_paragraph("本章概述核心结论，关键数据（见表1-1）。")
    doc.add_paragraph("表1-1 执行摘要关键数据与来源")
    report = tmp_path / "report.docx"
    doc.save(report)

    pipeline = MarketProductionPipeline(tmp_path, None, scripts_dir=tmp_path)
    pipeline._ensure_table_references(report)

    body = "".join(p.text for p in Document(report).paragraphs)
    assert body.count("（见表1-1）") == 1  # 已存在则不改动
