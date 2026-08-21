"""Word renderer (P0 refactor): narrative-driven, diagram-design PNG figures.

The Word document is built from the SAME ResearchNarrative as the HTML
dashboard.  Figures are high-resolution PNGs captured from the same
diagram-design HTML files that feed the HTML inline SVG — there is no
second charting implementation.  Chapters appear only when their evidence
gate passed (dynamic structure, driven by conclusions and decision
questions).

Internal diagnostics (renderer fallbacks, QA entries) are written to
``publication_qa_report.json`` in the assets folder and never appear in the
user-facing document.
"""

from __future__ import annotations

import hashlib
from datetime import date
from pathlib import Path

from enterprise_energy_research.adapters.base import AdapterHealth, ArtifactResult
from enterprise_energy_research.artifacts.diagram_design_adapter import DiagramDesignAdapter, VisualRenderResult
from enterprise_energy_research.artifacts.image_publication import (
    PublicationImage,
    prepare_publication_images,
    write_image_publication_manifest,
)
from enterprise_energy_research.artifacts.narrative import NarrativeBuilder, ResearchNarrative, write_narrative
from enterprise_energy_research.artifacts.qa_report import (
    QAFinding,
    QAVisualEntry,
    new_qa_report,
    write_qa_report,
)
from enterprise_energy_research.artifacts.visual_policy import colors as theme_colors
from enterprise_energy_research.artifacts.visual_policy import word_policy
from enterprise_energy_research.artifacts.visuals import VisualSpec, write_visual_manifest
from enterprise_energy_research.domain.enums import ArtifactType
from enterprise_energy_research.domain.models import ArtifactBinding, FrozenResearchBundle
from enterprise_energy_research.research.synthesis import ResearchSynthesizer


class FrozenWordPublisher:
    """Formal Word report: consulting style, narrative-driven chapters,
    same-source diagram-design figures, QA strictly outside the document."""

    name = "word_document"
    artifact_type = ArtifactType.WORD

    def health(self) -> AdapterHealth:
        try:
            import docx  # noqa: F401
            return AdapterHealth(name=self.name, available=True, version="python-docx")
        except ImportError:
            return AdapterHealth(name=self.name, available=False, diagnostics=["python-docx is unavailable"])

    def publish(self, bundle: FrozenResearchBundle, binding: ArtifactBinding, output_path: Path) -> ArtifactResult:
        health = self.health()
        if not health.available:
            return ArtifactResult(adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type, status="failed", diagnostics=health.diagnostics)
        if binding.type != self.artifact_type:
            return ArtifactResult(adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type, status="failed", diagnostics=["Word publisher received a non-Word binding"])

        entity = self._canonical_entity(bundle)
        if entity is None:
            return ArtifactResult(adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type, status="failed", diagnostics=["Frozen bundle contains no enterprise entity"])
        synthesis = ResearchSynthesizer().synthesize(
            run_id=bundle.run_manifest.run_id,
            entity=entity,
            entities=bundle.entities,
            claims=bundle.claims,
            sources=bundle.sources,
            edges=bundle.edges,
            factories=bundle.factories,
            products=bundle.products,
            energy_profiles=bundle.energy_profiles,
            gaps=bundle.gaps,
            solutions=bundle.solutions,
        )
        narrative = NarrativeBuilder().build(bundle, synthesis)

        asset_root = output_path.parent / f"{output_path.stem}_assets"
        figures = asset_root / "figures"
        adapter = DiagramDesignAdapter()
        qa = new_qa_report(bundle.run_manifest.run_id, bundle.freeze.freeze_id, binding.artifact_id)

        render_results: dict[str, VisualRenderResult] = {}
        for spec in narrative.visuals:
            result = adapter.build_visual(spec, figures, destination="both", png_scale=3)
            render_results[spec.visual_id] = result
            outcome = "rendered" if result.status == "rendered" else result.status
            qa.record_visual(QAVisualEntry(
                visual_id=spec.visual_id, chapter_id=spec.chapter_id,
                outcome=outcome,  # type: ignore[arg-type]
                visual_type=result.visual_type,
                reason=result.fallback_reason or result.error,
                png_status=result.png_status,
            ))
            if result.status == "failed":
                qa.record_finding(QAFinding(
                    code="visual_render_failed", severity="error",
                    message=f"{spec.visual_id} could not be rendered; insight kept as prose",
                    record_ids=[spec.visual_id],
                ))
            elif result.status == "fallback_table":
                qa.record_finding(QAFinding(
                    code="visual_fallback_table", severity="warn",
                    message=f"{spec.visual_id} degraded to structured table: {result.fallback_reason}",
                    record_ids=[spec.visual_id],
                ))

        write_visual_manifest(narrative.visual_manifest(), asset_root / "visual_manifest.json")
        write_narrative(narrative, asset_root / "narrative.json")

        image_manifest = prepare_publication_images(
            bundle, binding, asset_root, extra_search_roots=[output_path.parent],
        )
        publication_images = {item.image_id: item for item in image_manifest.prepared_images}
        for image_id, reason in image_manifest.withheld_reasons.items():
            qa.record_finding(QAFinding(
                code="image_withheld", severity="info",
                message=f"{image_id} withheld from publication: {reason}",
                record_ids=[image_id],
            ))
        selected_word_image_ids = [item.image_id for item in image_manifest.prepared_images]
        image_manifest = image_manifest.model_copy(update={"artifact_selections": {"word": selected_word_image_ids}})
        write_image_publication_manifest(image_manifest, asset_root)

        self._build_document(
            bundle, binding, output_path, entity, synthesis, narrative,
            render_results, publication_images, asset_root,
        )
        write_qa_report(qa, asset_root / "publication_qa_report.json")

        digest = hashlib.sha256(output_path.read_bytes()).hexdigest()
        used_claims = sorted({claim.claim_id for claim in self._claims_used(narrative, bundle)})
        return ArtifactResult(
            adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type,
            path=output_path, content_sha256=digest, used_claim_ids=used_claims,
            used_image_ids=selected_word_image_ids, status="published",
            diagnostics=image_manifest.diagnostics,
        )

    @staticmethod
    def _canonical_entity(bundle: FrozenResearchBundle):
        return next(
            (item for item in bundle.entities if item.entity_id == bundle.run_manifest.canonical_entity_id),
            bundle.entities[0] if bundle.entities else None,
        )

    @staticmethod
    def _claims_used(narrative: ResearchNarrative, bundle: FrozenResearchBundle) -> list:
        ids: set[str] = set()
        for chapter in narrative.chapters:
            ids.update(chapter.claim_ids)
        for visual in narrative.visuals:
            ids.update(visual.source_claim_ids)
        return [claim for claim in bundle.claims if claim.claim_id in ids]

    # ── document assembly ──
    def _build_document(
        self,
        bundle: FrozenResearchBundle,
        binding: ArtifactBinding,
        output_path: Path,
        entity,
        synthesis,
        narrative: ResearchNarrative,
        render_results: dict[str, VisualRenderResult],
        publication_images: dict[str, PublicationImage],
        asset_root: Path,
    ) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Mm, Pt, RGBColor

        wp = word_policy()
        tc = theme_colors()
        figure_width = Cm(wp["maximum_figure_width_cm"])
        body_cjk = wp["body_cjk_font"]
        body_latin = wp["body_latin_font"]
        navy_hex = tc["navy"].lstrip("#")
        cool_gray_hex = tc["cool_gray"].lstrip("#")

        document = Document()
        section = document.sections[0]
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Mm(25.4)
        section.header_distance = section.footer_distance = Mm(12.5)

        styles = document.styles
        for name, size, color, before, after, cjk_font in [
            ("Normal", wp["body_size_pt"], "1B1F26", 0, 6, body_cjk),
            ("Heading 1", wp["heading_1_size_pt"], navy_hex, 18, 10, "Microsoft YaHei"),
            ("Heading 2", wp["heading_2_size_pt"], navy_hex, 14, 7, body_cjk),
            ("Heading 3", wp["heading_3_size_pt"], navy_hex, 10, 5, body_cjk),
        ]:
            style = styles[name]
            style.font.name = body_latin
            style._element.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)
            style._element.rPr.rFonts.set(qn("w:ascii"), body_latin)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), body_latin)
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.bold = name != "Normal"
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = name != "Normal"
            if name == "Normal":
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                style.paragraph_format.line_spacing = Pt(wp["line_spacing_pt"])
                style.paragraph_format.first_line_indent = Pt(wp["body_size_pt"] * wp["first_line_indent_characters"])
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif name == "Heading 1":
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

        header = section.header.paragraphs[0]
        header.text = "企业产业与能源合作智能调研"
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor.from_string(cool_gray_hex)
        today = date.today()
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run(f"数据来源：公开渠道已核验证据（详见附录来源清单） · {today:%Y-%m-%d} · ")
        footer.add_run("偏差说明：本报告基于公开信息编制，不构成投资建议。 · ")
        self._field(footer, "PAGE")
        # ── cover (consulting: navy accents, no decoration) ──
        cover_spacer = document.add_paragraph()
        cover_spacer.paragraph_format.space_after = Pt(72)
        kicker = document.add_paragraph("企业研究 · 产业与能源合作")
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kr = kicker.runs[0]
        kr.font.name = "Arial"
        kr.bold = True
        kr.font.size = Pt(11)
        kr.font.color.rgb = RGBColor.from_string(navy_hex)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(18)
        tr = title.add_run(entity.canonical_name)
        tr.bold = True
        tr.font.name = "Microsoft YaHei"
        tr.font.size = Pt(26)
        tr.font.color.rgb = RGBColor.from_string("1B1F26")
        subtitle = document.add_paragraph("企业产业与能源合作智能调研报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(26)
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].bold = True
        subtitle.runs[0].font.color.rgb = RGBColor.from_string(navy_hex)
        rule = document.add_paragraph()
        rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pr = rule._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:color"), navy_hex)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        rule.paragraph_format.space_after = Pt(30)
        report_no = f"EER-{today:%Y%m%d}-{bundle.freeze.freeze_id[-6:]}"
        for label, value in (
            ("报告编号", report_no),
            ("数据版本", bundle.freeze.freeze_id),
            ("数据截止", bundle.freeze.created_at.date().isoformat()),
            ("生成日期", f"{today:%Y}年{today.month}月{today.day}日"),
        ):
            meta = document.add_paragraph(f"{label}：{value}")
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.paragraph_format.space_before = Pt(4)
            meta.paragraph_format.space_after = Pt(4)
            for run in meta.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string(cool_gray_hex)
        document.add_page_break()

        document.add_heading("目录", level=1)
        toc = document.add_paragraph()
        self._field(toc, 'TOC \\o "1-3" \\h \\z \\u')
        document.add_page_break()

        document.add_heading("决策问题", level=1)
        for index, question in enumerate(narrative.decision_questions, start=1):
            document.add_paragraph(f"问题 {index}：{question}")
        document.add_page_break()

        # ── narrative-driven chapters ──
        for index, chapter in enumerate(narrative.chapters, start=1):
            document.add_heading(f"{index}. {chapter.title}", level=1)
            if chapter.thesis and chapter.thesis.strip():
                thesis = document.add_paragraph(chapter.thesis)
                thesis.paragraph_format.first_line_indent = Pt(0)
                for run in thesis.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(navy_hex)
            for paragraph in chapter.content:
                document.add_paragraph(paragraph)
            if chapter.table_rows:
                self._add_structured_table(document, chapter.table_rows, f"表 {index}-1 {chapter.title}")
            for figure_no, visual_id in enumerate(chapter.visual_ids, start=1):
                result = render_results.get(visual_id)
                spec = next((item for item in narrative.visuals if item.visual_id == visual_id), None)
                if result is None or spec is None:
                    continue
                self._add_visual(document, spec, result, f"{index}-{figure_no}", figure_width, asset_root)
            image_counter = 0
            for image_id in chapter.image_ids:
                publication = publication_images.get(image_id)
                if publication is None:
                    continue
                image_counter += 1
                self._add_evidence_image(document, publication, asset_root, f"{index}-P{image_counter}", figure_width)

        # ── appendices ──
        document.add_page_break()
        document.add_heading("附录 A：术语与口径", level=1)
        document.add_paragraph(
            "本报告中的事实均来自公开渠道并经过核验；分析推断、待确认事项分别标注。"
            "产能、收入、能耗等数值以原始披露的时间、范围、单位和限定条件为准，"
            "不进行行业均值替代。"
        )
        document.add_heading("附录 B：来源清单", level=1)
        for source in bundle.sources:
            document.add_paragraph(f"{source.source_title or source.source_domain}｜{source.canonical_url}")
        document.add_heading("附录 C：图片来源", level=1)
        if publication_images:
            for publication in publication_images.values():
                document.add_paragraph(f"{publication.caption}｜{publication.source_page_url}")
        else:
            document.add_paragraph("本次研究未包含满足核验标准的实体图片。")
        document.add_heading("附录 D：待尽调事项", level=1)
        for gap in bundle.gaps:
            if gap.importance == "minor":
                continue
            document.add_heading(gap.field_name, level=2)
            document.add_paragraph(gap.next_action)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    # ── figure/table helpers ──
    @classmethod
    def _add_visual(
        cls,
        document,
        spec: VisualSpec,
        result: VisualRenderResult,
        figure_no: str,
        width,
        asset_root: Path,
    ) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Cm, Pt, RGBColor

        lead = document.add_paragraph(
            f"结论：{spec.business_thesis} 对应问题：{spec.decision_question}（见图 {figure_no}）。"
        )
        lead.paragraph_format.keep_with_next = True
        lead.paragraph_format.first_line_indent = Pt(0)
        for run in lead.runs:
            run.bold = True
        if result.png_path is not None and result.png_path.is_file():
            cls._insert_picture(document, result.png_path, width)
        elif spec.items:
            # PNG pipeline unavailable: structured table from the SAME spec data.
            rows = [
                {"指标": item.label, "数值": f"{item.value or ''} {item.unit or ''}".strip(),
                 "期间": item.period or "", "说明": item.note or ""}
                for item in spec.items
            ]
            cls._add_structured_table(document, rows, f"表 {figure_no} {spec.title}")
        else:
            return  # insight already kept as prose; QA recorded the failure
        caption = document.add_paragraph(f"图 {figure_no} {spec.title}")
        cls._format_caption(caption)
        source = document.add_paragraph(spec.source_note or "数据来源：公开信息（详见附录来源清单）。")
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        source.paragraph_format.space_before = Pt(0)
        source.paragraph_format.space_after = Pt(8)
        for run in source.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(74, 85, 104)

    @classmethod
    def _add_evidence_image(cls, document, image: PublicationImage, asset_root: Path, figure_no: str, width) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Cm, Pt, RGBColor

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        ratio = image.width / image.height if image.height else 1.0
        picture_width = Cm(13.8 if ratio >= 1.1 else 9.5)
        paragraph.add_run().add_picture(str(asset_root / image.publication_path), width=picture_width)
        caption = document.add_paragraph(f"图 {figure_no} {image.caption}")
        cls._format_caption(caption)
        source = document.add_paragraph(image.source_note)
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        source.paragraph_format.space_before = Pt(0)
        source.paragraph_format.space_after = Pt(10)
        for run in source.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(74, 85, 104)

    @staticmethod
    def _insert_picture(document, png_path: Path, width) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Pt

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run().add_picture(str(png_path), width=width)

    @classmethod
    def _add_structured_table(cls, document, rows: list[dict], caption: str) -> None:
        if not rows:
            return
        table_caption = document.add_paragraph(caption)
        cls._format_caption(table_caption)
        columns = list(rows[0].keys())
        table = document.add_table(rows=1, cols=len(columns))
        table.autofit = False
        for cell, text in zip(table.rows[0].cells, columns):
            cell.text = str(text)
        for row in rows:
            cells = table.add_row().cells
            for cell, column in zip(cells, columns):
                cell.text = str(row.get(column) or "")
        cell_width = 9360 // max(len(columns), 1)
        cls._table_geometry(table, [cell_width] * len(columns))
        cls._style_three_line_table(table)

    @staticmethod
    def _format_caption(paragraph) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Pt

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

    @classmethod
    def _style_three_line_table(cls, table) -> None:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        tc = theme_colors()
        navy_hex = tc["navy"].lstrip("#")
        pale_hex = tc["canvas"].lstrip("#")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_pr = table._tbl.tblPr
        existing = tbl_pr.find(qn("w:tblBorders"))
        if existing is not None:
            tbl_pr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge, value, size, color in (
            ("top", "single", "12", "000000"), ("bottom", "single", "12", "000000"),
            ("left", "nil", "0", "FFFFFF"), ("right", "nil", "0", "FFFFFF"),
            ("insideH", "nil", "0", "FFFFFF"), ("insideV", "nil", "0", "FFFFFF"),
        ):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), value)
            node.set(qn("w:sz"), size)
            node.set(qn("w:color"), color)
            borders.append(node)
        tbl_pr.append(borders)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                if row_index == 0:
                    shade = OxmlElement("w:shd")
                    shade.set(qn("w:fill"), pale_hex)
                    tc_pr.append(shade)
                    cell_border = OxmlElement("w:tcBorders")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "8")
                    bottom.set(qn("w:color"), navy_hex)
                    cell_border.append(bottom)
                    tc_pr.append(cell_border)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.left_indent = Pt(0)
                    paragraph.paragraph_format.right_indent = Pt(0)
                    paragraph.paragraph_format.keep_together = True
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(word_policy()["table_size_pt"])
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                        if row_index == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor.from_string(navy_hex)

    @staticmethod
    def _field(paragraph, instruction: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        text.text = "更新域以显示"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for node in (begin, instr, separate, text, end):
            run._r.append(node)

    @staticmethod
    def _table_geometry(table, widths: list[int]) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), str(sum(widths)))
        tbl_w.set(qn("w:type"), "dxa")
        if tbl_w.getparent() is None:
            tbl_pr.append(tbl_w)
        existing_layout = tbl_pr.find(qn("w:tblLayout"))
        if existing_layout is not None:
            tbl_pr.remove(existing_layout)
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), "fixed")
        tbl_pr.append(tbl_layout)
        existing_ind = tbl_pr.find(qn("w:tblInd"))
        if existing_ind is not None:
            tbl_pr.remove(existing_ind)
        tbl_ind = OxmlElement("w:tblInd")
        tbl_ind.set(qn("w:w"), "0")
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_ind)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_w.set(qn("w:w"), str(width))
                tc_w.set(qn("w:type"), "dxa")
