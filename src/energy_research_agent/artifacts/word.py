"""Word renderer (P0 refactor): narrative-driven, diagram-design PNG figures.

The Word document is built from the SAME ResearchNarrative as the HTML
dashboard.  Figures are high-resolution PNGs captured from the same
diagram-design HTML files that feed the HTML inline SVG — there is no
second charting implementation.  Chapters appear only when their evidence
gate passed (dynamic structure, driven by conclusions and decision
questions).

Internal diagnostics (renderer fallbacks, QA entries) are written to
``publication_qa_report.json`` in the assets folder and never appear in the
user-facing document.
"""

from __future__ import annotations

import hashlib
from datetime import date
from pathlib import Path

from energy_research_agent.adapters.base import AdapterHealth, ArtifactResult
from energy_research_agent.artifacts.diagram_design_adapter import DiagramDesignAdapter, VisualRenderResult
from energy_research_agent.artifacts.image_publication import (
    PublicationImage,
    prepare_publication_images,
    write_image_publication_manifest,
)
from energy_research_agent.artifacts.narrative import NarrativeBuilder, ResearchNarrative, write_narrative
from energy_research_agent.artifacts.publication_boilerplate import PublicationBoilerplateFilter
from energy_research_agent.artifacts.toc import heading_levels, materialize_toc
from energy_research_agent.artifacts.qa_report import (
    QAFinding,
    QAVisualEntry,
    downgrade_conditional_findings,
    new_qa_report,
    write_qa_report,
)
from energy_research_agent.artifacts.visual_policy import colors as theme_colors
from energy_research_agent.artifacts.visual_policy import word_policy
from energy_research_agent.artifacts.visuals import VisualSpec, write_visual_manifest
from energy_research_agent.domain.enums import ArtifactType, VerificationStatus
from energy_research_agent.domain.models import ArtifactBinding, FrozenResearchBundle
from energy_research_agent.research.synthesis import ResearchSynthesizer
from energy_research_agent.validation.consulting_narrative import (
    ConsultingNarrativeValidator, PublicationVisibleTextValidator, TOCValidator,
    VisualSemanticValidator, write_consulting_validation,
)
from energy_research_agent.validation.publication_quality import (
    DecisionIntelligenceValidator,
    ProductImageCoverageValidator,
    PublicationBoilerplateValidator,
    ResearchValueValidator,
)


class FrozenWordPublisher:
    """Formal Word report: consulting style, narrative-driven chapters,
    same-source diagram-design figures, QA strictly outside the document."""

    name = "word_document"
    artifact_type = ArtifactType.WORD

    def health(self) -> AdapterHealth:
        try:
            import docx  # noqa: F401
            return AdapterHealth(name=self.name, available=True, version="python-docx")
        except ImportError:
            return AdapterHealth(name=self.name, available=False, diagnostics=["python-docx is unavailable"])

    def publish(self, bundle: FrozenResearchBundle, binding: ArtifactBinding, output_path: Path) -> ArtifactResult:
        health = self.health()
        if not health.available:
            return ArtifactResult(adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type, status="failed", diagnostics=health.diagnostics)
        if binding.type != self.artifact_type:
            return ArtifactResult(adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type, status="failed", diagnostics=["Word publisher received a non-Word binding"])

        entity = self._canonical_entity(bundle)
        if entity is None:
            return ArtifactResult(adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type, status="failed", diagnostics=["Frozen bundle contains no enterprise entity"])
        synthesis = ResearchSynthesizer().synthesize(
            run_id=bundle.run_manifest.run_id,
            entity=entity,
            entities=bundle.entities,
            claims=bundle.claims,
            sources=bundle.sources,
            edges=bundle.edges,
            factories=bundle.factories,
            products=bundle.products,
            energy_profiles=bundle.energy_profiles,
            gaps=bundle.gaps,
            solutions=bundle.solutions,
        )
        narrative = NarrativeBuilder().build(bundle, synthesis)
        narrative = ResearchNarrative.model_validate(
            PublicationBoilerplateFilter().filter_value(narrative.model_dump(mode="json"))
        )

        asset_root = output_path.parent / f"{output_path.stem}_assets"
        figures = asset_root / "figures"
        adapter = DiagramDesignAdapter()
        qa = new_qa_report(bundle.run_manifest.run_id, bundle.freeze.freeze_id, binding.artifact_id)
        fixture_mode = bundle.run_manifest.model_gateway.get("mode") in {
            "fixture", "recorded-fixture", "recorded-fixture-only",
        }
        narrative_validation = ConsultingNarrativeValidator().validate(
            narrative, enforce_length=not fixture_mode,
        )
        write_consulting_validation(narrative_validation, asset_root / "consulting_narrative_validation.json")
        for check in narrative_validation.checks:
            if check.status == "FAIL":
                qa.record_finding(QAFinding(
                    code=check.code, severity="error", message=check.message,
                ))
        for check in [
            *PublicationBoilerplateValidator().validate(narrative),
            *ResearchValueValidator().validate(narrative, bundle),
            *ProductImageCoverageValidator().validate(narrative),
            *DecisionIntelligenceValidator().validate(narrative, bundle),
        ]:
            if check.status == "FAIL":
                qa.record_finding(QAFinding(code=check.code, severity="error", message=check.message))
            elif check.status == "WARN":
                qa.record_finding(QAFinding(code=check.code, severity="warn", message=check.message))

        render_results: dict[str, VisualRenderResult] = {}
        for spec in narrative.visuals:
            for semantic_error in VisualSemanticValidator().validate(spec, bundle):
                qa.record_finding(QAFinding(
                    code="visual_semantic_violation", severity="error", message=semantic_error,
                    record_ids=[spec.visual_id],
                ))
            result = adapter.build_visual(spec, figures, destination="both", png_scale=3)
            render_results[spec.visual_id] = result
            outcome = (
                "fallback_table"
                if result.status == "rendered" and result.png_status != "ok"
                else "rendered" if result.status == "rendered" else result.status
            )
            qa.record_visual(QAVisualEntry(
                visual_id=spec.visual_id, chapter_id=spec.chapter_id,
                outcome=outcome,  # type: ignore[arg-type]
                visual_type=result.visual_type,
                reason=result.fallback_reason or result.error,
                png_status=result.png_status,
            ))
            if result.status == "failed":
                qa.record_finding(QAFinding(
                    code="visual_render_failed", severity="error",
                    message=f"{spec.visual_id} could not be rendered; insight kept as prose",
                    record_ids=[spec.visual_id],
                ))
            elif result.status == "rendered" and result.png_status != "ok":
                qa.record_finding(QAFinding(
                    code="word_visual_degraded_to_table", severity="warn",
                    message=(f"{spec.visual_id} has no Word-ready PNG and was "
                             "published as a structured table from the same VisualSpec data"),
                    record_ids=[spec.visual_id],
                ))
            elif result.status == "fallback_table":
                qa.record_finding(QAFinding(
                    code="visual_fallback_table", severity="warn",
                    message=f"{spec.visual_id} degraded to structured table: {result.fallback_reason}",
                    record_ids=[spec.visual_id],
                ))

        write_visual_manifest(narrative.visual_manifest(), asset_root / "visual_manifest.json")
        write_narrative(narrative, asset_root / "narrative.json")

        image_manifest = prepare_publication_images(
            bundle, binding, asset_root, extra_search_roots=[output_path.parent],
        )
        publication_images = {item.image_id: item for item in image_manifest.prepared_images}
        for image_id, reason in image_manifest.withheld_reasons.items():
            qa.record_finding(QAFinding(
                code="image_withheld", severity="info",
                message=f"{image_id} withheld from publication: {reason}",
                record_ids=[image_id],
            ))
        selected_word_image_ids = [item.image_id for item in image_manifest.prepared_images]
        image_manifest = image_manifest.model_copy(update={"artifact_selections": {"word": selected_word_image_ids}})
        write_image_publication_manifest(image_manifest, asset_root)

        self._build_document(
            bundle, binding, output_path, entity, synthesis, narrative,
            render_results, publication_images, asset_root,
        )
        # A real TOC field with an empty cached result looks like a missing
        # directory outside desktop Word.  Seed visible Heading 1/2 entries;
        # final office-render QA refreshes their page-number slots.
        toc_headings = heading_levels(output_path)
        materialize_toc(output_path, [(heading, 0, level) for heading, level in toc_headings])
        self._render_pdf_sidecar(output_path, qa)
        visible_validator = PublicationVisibleTextValidator()
        for message in [*visible_validator.validate_text(visible_validator.extract_docx(output_path)), *TOCValidator().validate(output_path)]:
            qa.record_finding(QAFinding(code="word_visible_text_or_toc", severity="error", message=message))
        verified_products = [item for item in bundle.products if item.verification_status == VerificationStatus.VERIFIED]
        verified_product_ids = {item.product_id for item in verified_products}
        image_backed_product_ids = {
            item.product_id for item in image_manifest.prepared_images
            if item.product_id in verified_product_ids
        }
        # Official product photography bound to the subject entity (verified +
        # pixel-checked) also counts toward product image coverage when no
        # per-product binding exists.
        entity_bound_product_photos = {
            item.image_id for item in image_manifest.prepared_images
            if not item.product_id and item.target_entity_id
            and item.target_entity_type == "product" and item.visual_verified
        }
        coverage = max(len(image_backed_product_ids), len(entity_bound_product_photos))
        if len(verified_products) >= 5 and coverage < 5:
            qa.record_finding(QAFinding(
                code="word_product_image_gate", severity="error",
                message="正式 Word 报告至少需要 5 个图片、参数和场景完整绑定的重点产品。",
            ))
        downgrade_conditional_findings(qa, bundle)
        write_qa_report(qa, asset_root / "publication_qa_report.json")

        digest = hashlib.sha256(output_path.read_bytes()).hexdigest()
        used_claims = sorted({claim.claim_id for claim in self._claims_used(narrative, bundle)})
        return ArtifactResult(
            adapter=self.name, artifact_id=binding.artifact_id, artifact_type=binding.type,
            path=output_path, content_sha256=digest, used_claim_ids=used_claims,
            used_image_ids=selected_word_image_ids, status="published" if qa.status != "fail" else "failed",
            diagnostics=[*image_manifest.diagnostics, *([] if qa.status != "fail" else ["publication_qa_failed"])],
        )

    @staticmethod
    def _canonical_entity(bundle: FrozenResearchBundle):
        return next(
            (item for item in bundle.entities if item.entity_id == bundle.run_manifest.canonical_entity_id),
            bundle.entities[0] if bundle.entities else None,
        )

    @staticmethod
    def _render_pdf_sidecar(output_path: Path, qa) -> None:
        """Render the same-name PDF the release word-depth gate consumes.

        ``ArtifactConsistencyAuditor`` calls ``inspect_word_depth(path,
        rendered_pdf=path.with_suffix('.pdf'))`` outside fixture mode and the
        gate hard-blocks when the rendered PDF is missing.  LibreOffice
        headless is the only office renderer available in the container, so
        its absence must surface as a QA finding instead of a silent gate
        failure at audit time.
        """
        import shutil
        import subprocess

        soffice = shutil.which("soffice") or shutil.which("libreoffice")
        if not soffice:
            qa.record_finding(QAFinding(
                code="word_pdf_render_unavailable", severity="warn",
                message=(
                    "LibreOffice is unavailable; the formal-depth gate requires "
                    "a rendered PDF sibling next to the Word report"
                ),
            ))
            return
        target = output_path.with_suffix(".pdf")
        try:
            proc = subprocess.run(
                [soffice, "--headless", "--convert-to", "pdf",
                 "--outdir", str(output_path.parent), str(output_path)],
                capture_output=True, text=True, timeout=900,
            )
        except (OSError, subprocess.TimeoutExpired) as exc:
            qa.record_finding(QAFinding(
                code="word_pdf_render_failed", severity="warn",
                message=f"LibreOffice PDF render failed: {type(exc).__name__}: {exc}",
            ))
            return
        if proc.returncode != 0 or not target.is_file():
            tail = ((proc.stderr or proc.stdout) or "").strip()[-300:]
            qa.record_finding(QAFinding(
                code="word_pdf_render_failed", severity="warn",
                message=f"LibreOffice PDF render failed (rc={proc.returncode}): {tail}",
            ))

    @staticmethod
    def _claims_used(narrative: ResearchNarrative, bundle: FrozenResearchBundle) -> list:
        ids: set[str] = set()
        for chapter in narrative.chapters:
            ids.update(chapter.claim_ids)
        for visual in narrative.visuals:
            ids.update(visual.source_claim_ids)
        return [claim for claim in bundle.claims if claim.claim_id in ids]

    # ── document assembly ──
    def _build_document(
        self,
        bundle: FrozenResearchBundle,
        binding: ArtifactBinding,
        output_path: Path,
        entity,
        synthesis,
        narrative: ResearchNarrative,
        render_results: dict[str, VisualRenderResult],
        publication_images: dict[str, PublicationImage],
        asset_root: Path,
    ) -> None:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Cm, Mm, Pt, RGBColor

        wp = word_policy()
        tc = theme_colors()
        figure_width = Cm(wp["maximum_figure_width_cm"])
        body_cjk = wp["body_cjk_font"]
        body_latin = wp["body_latin_font"]
        navy_hex = tc["navy"].lstrip("#")
        cool_gray_hex = tc["cool_gray"].lstrip("#")

        document = Document()
        section = document.sections[0]
        section.page_width, section.page_height = Mm(210), Mm(297)
        section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Mm(25.4)
        section.header_distance = section.footer_distance = Mm(12.5)

        conditional_scope = bundle.run_manifest.research_scope or {}
        if conditional_scope.get("publication_mode") == "conditional":
            notice = document.add_paragraph()
            notice.paragraph_format.first_line_indent = Pt(0)
            notice.paragraph_format.space_after = Pt(12)
            notice_run = notice.add_run(
                "【条件发布】本报告基于已核验的公开证据编制，但存在未能通过公开渠道补齐的"
                "覆盖缺口（详见『附录 D：待尽调事项』与 formal_publication_eligibility.json）。"
                "相关结论受此证据边界限制，正式对外发布前需完成补充尽调。"
            )
            notice_run.font.bold = True
            notice_run.font.size = Pt(11)

        styles = document.styles
        for name, size, color, before, after, cjk_font in [
            ("Normal", wp["body_size_pt"], "1B1F26", 0, 6, body_cjk),
            ("Heading 1", wp["heading_1_size_pt"], navy_hex, 18, 12, "Microsoft YaHei"),
            ("Heading 2", wp["heading_2_size_pt"], navy_hex, 14, 7, "FangSong"),
            ("Heading 3", wp["heading_3_size_pt"], navy_hex, 10, 5, body_cjk),
        ]:
            style = styles[name]
            style.font.name = body_latin
            style._element.rPr.rFonts.set(qn("w:eastAsia"), cjk_font)
            style._element.rPr.rFonts.set(qn("w:ascii"), body_latin)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), body_latin)
            style.font.size = Pt(size)
            style.font.color.rgb = RGBColor.from_string(color)
            style.font.bold = name != "Normal"
            style.paragraph_format.space_before = Pt(before)
            style.paragraph_format.space_after = Pt(after)
            style.paragraph_format.keep_with_next = name != "Normal"
            if name == "Normal":
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                style.paragraph_format.line_spacing = Pt(wp["line_spacing_pt"])
                style.paragraph_format.first_line_indent = Pt(wp["body_size_pt"] * wp["first_line_indent_characters"])
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif name == "Heading 1":
                style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                style.paragraph_format.line_spacing = Pt(30)

        # TOC 1/2/3 styles must exist with LEFT alignment.  Word/LibreOffice
        # materialize the TOC field through these styles; without them the
        # entries inherit Normal (JUSTIFY), which distributes CJK glyphs
        # ("执 行 摘 要 与 决 策 建 议").  Each level gets a right tab stop
        # with a dot leader at the content width.
        from docx.enum.style import WD_STYLE_TYPE
        from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
        # Keep the executive TOC compact enough that a 9-chapter report plus
        # appendices does not spill one orphan entry onto a nearly blank page.
        for toc_name, size in (("TOC 1", 10), ("TOC 2", 9.5), ("TOC 3", 9)):
            try:
                toc_style = styles.add_style(toc_name, WD_STYLE_TYPE.PARAGRAPH)
            except ValueError:
                toc_style = styles[toc_name]
            toc_style.font.name = body_latin
            toc_style._element.rPr.rFonts.set(qn("w:eastAsia"), body_cjk)
            toc_style.font.size = Pt(size)
            toc_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            toc_style.paragraph_format.first_line_indent = Pt(0)
            toc_style.paragraph_format.left_indent = Pt(0)
            toc_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            toc_style.paragraph_format.tab_stops.add_tab_stop(Cm(15.9), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)

        # Make TOC/page fields refresh when Word or LibreOffice opens the file.
        settings = document.settings.element
        update_fields = OxmlElement("w:updateFields")
        update_fields.set(qn("w:val"), "true")
        settings.append(update_fields)

        header = section.header.paragraphs[0]
        header.text = "企业产业与能源合作研究报告"
        header.runs[0].font.size = Pt(9)
        header.runs[0].font.color.rgb = RGBColor.from_string(cool_gray_hex)
        today = date.today()
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer.add_run(f"数据来源：公开渠道已核验证据（详见附录来源清单） · {today:%Y-%m-%d} · ")
        footer.add_run("偏差说明：本报告基于公开信息编制，不构成投资建议。 · ")
        self._field(footer, "PAGE")
        # ── cover (consulting: navy accents, no decoration) ──
        cover_spacer = document.add_paragraph()
        cover_spacer.paragraph_format.space_after = Pt(72)
        kicker = document.add_paragraph("企业研究 · 产业与能源合作")
        kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
        kr = kicker.runs[0]
        kr.font.name = "Arial"
        kr.bold = True
        kr.font.size = Pt(11)
        kr.font.color.rgb = RGBColor.from_string(navy_hex)
        title = document.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(18)
        tr = title.add_run(entity.canonical_name)
        tr.bold = True
        tr.font.name = "Microsoft YaHei"
        tr.font.size = Pt(26)
        tr.font.color.rgb = RGBColor.from_string("1B1F26")
        subtitle = document.add_paragraph("企业产业与能源合作战略研究报告")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_after = Pt(26)
        subtitle.runs[0].font.size = Pt(16)
        subtitle.runs[0].bold = True
        subtitle.runs[0].font.color.rgb = RGBColor.from_string(navy_hex)
        rule = document.add_paragraph()
        rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_pr = rule._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:color"), navy_hex)
        p_bdr.append(bottom)
        p_pr.append(p_bdr)
        rule.paragraph_format.space_after = Pt(30)
        report_no = f"ERA-{today:%Y%m%d}-{bundle.freeze.freeze_id[-6:]}"
        for label, value in (
            ("报告编号", report_no),
            ("数据版本", bundle.freeze.freeze_id),
            ("数据截止", bundle.freeze.created_at.date().isoformat()),
            ("生成日期", f"{today:%Y}年{today.month}月{today.day}日"),
        ):
            meta = document.add_paragraph(f"{label}：{value}")
            meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
            meta.paragraph_format.space_before = Pt(4)
            meta.paragraph_format.space_after = Pt(4)
            for run in meta.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor.from_string(cool_gray_hex)
        document.add_page_break()

        document.add_heading("目录", level=1)
        toc = document.add_paragraph()
        self._field(toc, 'TOC \\o "1-2" \\h \\z \\u')
        document.add_page_break()

        # ── narrative-driven chapters ──
        for index, chapter in enumerate(narrative.chapters, start=1):
            document.add_heading(f"{index}. {chapter.title}", level=1)
            if chapter.assertion_title.strip():
                document.add_heading(chapter.assertion_title, level=2)
            if chapter.executive_takeaway.strip():
                takeaway = document.add_paragraph(chapter.executive_takeaway)
                takeaway.paragraph_format.first_line_indent = Pt(0)
                takeaway.paragraph_format.left_indent = Pt(10)
                for run in takeaway.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(navy_hex)
            analysis_budget = {
                "executive_summary": 4, "operations": 6, "products": 4,
                "factories": 4, "energy_profile": 3, "opportunities": 2,
                "action_plan": 2, "risks_evidence": 0,
            }.get(chapter.kind, 3)
            for paragraph in [*chapter.context_paragraphs, *chapter.analysis_paragraphs[:analysis_budget]]:
                document.add_paragraph(paragraph)
            if chapter.table_rows:
                shown_rows = chapter.table_rows
                caption = f"表 {index}-1 {chapter.title}"
                if chapter.chapter_id == "products" and len(shown_rows) > 8:
                    shown_rows = shown_rows[:8]
                    caption = f"表 {index}-1 重点产品（完整产品清单见附录 F）"
                self._add_structured_table(document, shown_rows, caption)
            for figure_no, visual_id in enumerate(chapter.visual_ids, start=1):
                result = render_results.get(visual_id)
                spec = next((item for item in narrative.visuals if item.visual_id == visual_id), None)
                if result is None or spec is None:
                    continue
                self._add_visual(document, spec, result, f"{index}-{figure_no}", figure_width, asset_root)
            showcased_image_ids: set[str] = set()
            if chapter.chapter_id == "products":
                showcased_image_ids = self._add_product_showcase(
                    document, bundle, narrative, publication_images, asset_root, index,
                )
            image_counter = 0
            for image_id in chapter.image_ids:
                if image_id in showcased_image_ids:
                    continue
                publication = publication_images.get(image_id)
                if publication is None:
                    continue
                image_counter += 1
                self._add_evidence_image(document, publication, asset_root, f"{index}-P{image_counter}", figure_width)
            self._add_statement_list(document, "业务含义", chapter.implications)
            self._add_statement_list(document, "建议", chapter.recommendations)
            self._add_statement_list(document, "可能改变判断的事项", chapter.counter_evidence)
            constraints = [*chapter.limitations]
            if chapter.kind == "risks_evidence":
                constraints = [*chapter.analysis_paragraphs, *constraints]
            self._add_statement_list(document, "关键约束与待确认", constraints)
            self._add_statement_list(document, "行动项", chapter.action_items)

        # ── appendices ──
        # Chapters may reference only a subset of verified images; the appendix
        # gallery below embeds the remainder so every valid image is published.
        chapter_bound_image_ids = {
            image_id for chapter in narrative.chapters for image_id in chapter.image_ids
        }
        document.add_page_break()
        document.add_heading("附录 A：术语与口径", level=1)
        document.add_paragraph(
            "本报告中的事实均来自公开渠道并经过核验；分析推断、待确认事项分别标注。"
            "产能、收入、能耗等数值以原始披露的时间、范围、单位和限定条件为准，"
            "不进行行业均值替代。"
        )
        document.add_heading("附录 B：来源清单", level=1)
        self._add_structured_table(document, narrative.appendices.source_ledger, "表 B-1 来源清单")
        document.add_heading("附录 C：图片来源", level=1)
        if publication_images:
            for publication in publication_images.values():
                document.add_paragraph(f"{publication.caption}｜{publication.source_page_url}")
        else:
            document.add_paragraph("本次研究未包含满足核验标准的实体图片。")
        document.add_heading("附录 G：图片证据全集", level=1)
        document.add_paragraph(
            "以下为全部通过像素级视觉核验的图片证据（正文章节未收录部分在此补全，未重复展示）。"
        )
        gallery_counter = 0
        for image_id, publication in publication_images.items():
            if image_id in chapter_bound_image_ids:
                continue
            gallery_counter += 1
            self._add_evidence_image(document, publication, asset_root, f"附-{gallery_counter}", figure_width)
        if gallery_counter == 0 and publication_images:
            document.add_paragraph("全部核验图片均已在正文章节展示。")
        elif gallery_counter == 0:
            document.add_paragraph("本次研究未包含满足核验标准的实体图片。")
        document.add_heading("附录 D：待尽调事项", level=1)
        for item in narrative.appendices.due_diligence:
            # Item-level entries remain navigable but stay outside the 1–2
            # level executive TOC; otherwise a large due-diligence registry
            # can turn the contents section into dozens of unusable pages.
            document.add_heading(item.item, level=3)
            document.add_paragraph(item.why_it_matters)
            document.add_paragraph(f"建议材料：{'、'.join(item.requested_materials)}；获取时点：{item.timing}；是否阻断决策：{'是' if item.decision_blocker else '否'}")
        document.add_heading("附录 E：生产基地清单", level=1)
        self._add_structured_table(document, narrative.appendices.factory_ledger, "表 E-1 生产基地清单")
        document.add_heading("附录 F：产品清单", level=1)
        self._add_structured_table(document, narrative.appendices.product_ledger, "表 F-1 产品清单")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    # ── figure/table helpers ──
    @classmethod
    def _add_product_showcase(
        cls,
        document,
        bundle: FrozenResearchBundle,
        narrative: ResearchNarrative,
        publication_images: dict[str, PublicationImage],
        asset_root: Path,
        chapter_no: int,
    ) -> set[str]:
        """Pair 4—8 key products with official imagery, parameters and uses."""
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Cm, Pt, RGBColor

        rows = []
        prepared_by_product: dict[str, tuple[str, PublicationImage]] = {}
        for prepared_id, prepared in publication_images.items():
            if prepared.product_id and prepared.product_id not in prepared_by_product:
                prepared_by_product[prepared.product_id] = (prepared_id, prepared)
        for product in bundle.products:
            if product.verification_status != VerificationStatus.VERIFIED:
                continue
            image_id = narrative.product_images.get(product.product_id) or product.image_id
            publication = publication_images.get(image_id or "")
            if publication is None and product.product_id in prepared_by_product:
                image_id, publication = prepared_by_product[product.product_id]
            if publication is not None:
                rows.append((product, image_id, publication))
        rows = rows[:8]
        if not rows:
            return set()

        document.add_heading("重点产品图谱", level=2)
        used: set[str] = set()
        for offset, (product, image_id, publication) in enumerate(rows, start=1):
            used.add(image_id)
            heading = document.add_heading(product.name, level=3)
            heading.paragraph_format.keep_with_next = True
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.keep_together = True
            # Inline pictures inherit Normal's fixed 22 pt line height unless
            # overridden.  Word/LibreOffice then clips the picture to a thin
            # strip even though the binary itself is high resolution.  Every
            # picture paragraph must use automatic single-line height.
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            ratio = publication.width / publication.height if publication.height else 1.0
            paragraph.add_run().add_picture(
                str(asset_root / publication.publication_path),
                width=Cm(12.8 if ratio >= 1.1 else 8.8),
            )
            meta_parts = [
                f"产品族：{product.category or '未披露'}",
                f"系列/型号：{' / '.join(item for item in (product.series, product.model) if item) or '未披露'}",
                f"应用场景：{'、'.join(product.applications[:6]) or '未披露'}",
            ]
            meta = document.add_paragraph("；".join(meta_parts))
            meta.paragraph_format.first_line_indent = Pt(0)
            meta.paragraph_format.keep_together = True
            parameters = "；".join(
                f"{item.name}：{item.value if item.value is not None else ''}{item.unit or ''}"
                for item in product.parameters[:8]
            ) or "公开参数：未披露"
            parameter_paragraph = document.add_paragraph(f"关键参数：{parameters}")
            parameter_paragraph.paragraph_format.first_line_indent = Pt(0)
            source = document.add_paragraph(
                f"图 {chapter_no}-P{offset} {product.name}｜图片来源：{publication.source_page_url}｜产品来源编号：{'、'.join(product.source_ids) or '详见附录 F'}"
            )
            source.alignment = WD_ALIGN_PARAGRAPH.CENTER
            source.paragraph_format.first_line_indent = Pt(0)
            source.paragraph_format.space_after = Pt(10)
            for run in source.runs:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(74, 85, 104)
        return used

    @classmethod
    def _add_visual(
        cls,
        document,
        spec: VisualSpec,
        result: VisualRenderResult,
        figure_no: str,
        width,
        asset_root: Path,
    ) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Cm, Pt, RGBColor

        if result.png_path is not None and result.png_path.is_file():
            cls._insert_picture(document, result.png_path, width)
        elif spec.items:
            # PNG pipeline unavailable: structured table from the SAME spec data.
            rows = [
                {"指标": item.label, "数值": f"{item.value or ''} {item.unit or ''}".strip(),
                 "期间": item.period or "", "说明": item.note or ""}
                for item in spec.items
            ]
            cls._add_structured_table(document, rows, f"表 {figure_no} {spec.title}")
        else:
            return  # insight already kept as prose; QA recorded the failure
        caption = document.add_paragraph(f"图 {figure_no} {spec.title}")
        cls._format_caption(caption)
        source = document.add_paragraph(spec.source_note or "数据来源：公开信息（详见附录来源清单）。")
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        source.paragraph_format.space_before = Pt(0)
        source.paragraph_format.space_after = Pt(8)
        for run in source.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(74, 85, 104)
        so_what = document.add_paragraph(spec.business_thesis)
        so_what.paragraph_format.first_line_indent = Pt(0)
        so_what.paragraph_format.space_after = Pt(10)
        for run in so_what.runs:
            run.font.color.rgb = RGBColor(27, 54, 93)

    @staticmethod
    def _add_statement_list(document, label: str, items: list[str]) -> None:
        if not items:
            return
        from docx.shared import Pt

        heading = document.add_paragraph(label)
        heading.paragraph_format.first_line_indent = Pt(0)
        heading.paragraph_format.space_before = Pt(8)
        heading.paragraph_format.space_after = Pt(3)
        for run in heading.runs:
            run.bold = True
        for item in items:
            paragraph = document.add_paragraph(item, style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)

    @classmethod
    def _add_evidence_image(cls, document, image: PublicationImage, asset_root: Path, figure_no: str, width) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Cm, Pt, RGBColor

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        ratio = image.width / image.height if image.height else 1.0
        picture_width = Cm(13.8 if ratio >= 1.1 else 9.5)
        paragraph.add_run().add_picture(str(asset_root / image.publication_path), width=picture_width)
        caption = document.add_paragraph(f"图 {figure_no} {image.caption}")
        cls._format_caption(caption)
        source = document.add_paragraph(image.source_note)
        source.alignment = WD_ALIGN_PARAGRAPH.CENTER
        source.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        source.paragraph_format.space_before = Pt(0)
        source.paragraph_format.space_after = Pt(10)
        for run in source.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(74, 85, 104)

    @staticmethod
    def _insert_picture(document, png_path: Path, width) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Pt

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(6)
        paragraph.add_run().add_picture(str(png_path), width=width)

    @classmethod
    def _add_structured_table(cls, document, rows: list[dict], caption: str) -> None:
        if not rows:
            return
        rows = cls._compact_table_rows(rows)
        table_caption = document.add_paragraph(caption)
        cls._format_caption(table_caption)
        columns = list(rows[0].keys())
        table = document.add_table(rows=1, cols=len(columns))
        table.autofit = False
        for cell, text in zip(table.rows[0].cells, columns):
            cell.text = str(text)
        for row in rows:
            cells = table.add_row().cells
            for cell, column in zip(cells, columns):
                cell.text = str(row.get(column) or "")
        cls._table_geometry(table, cls._table_widths(columns))
        cls._style_three_line_table(table)

    @staticmethod
    def _compact_table_rows(rows: list[dict]) -> list[dict]:
        """Keep Word tables readable without weakening the evidence appendix.

        The shared narrative deliberately carries publication-complete rows.
        A portrait Word page, however, cannot legibly render six prose-heavy
        columns.  Compact only the Word projection: the HTML/shared narrative
        remains unchanged, while the complete due-diligence registry is still
        printed item-by-item in Appendix D.
        """
        if not rows:
            return rows
        columns = set(rows[0])
        due_columns = {
            "当前尚不能判断的关键事项", "为什么重要", "影响判断",
            "建议获取资料", "获取时点", "是否阻断决策",
        }
        if due_columns.issubset(columns):
            unique: list[dict] = []
            seen: set[str] = set()
            ordered = sorted(rows, key=lambda row: row.get("是否阻断决策") != "是")
            for row in ordered:
                item = str(row.get("当前尚不能判断的关键事项") or "").strip()
                if not item or item in seen:
                    continue
                seen.add(item)
                unique.append({
                    "关键事项": item,
                    "重要性与影响": (
                        f"{row.get('为什么重要') or ''}；影响判断：{row.get('影响判断') or ''}"
                    ).strip("；"),
                    "建议获取资料": row.get("建议获取资料") or "",
                    "时点与门槛": (
                        f"{row.get('获取时点') or ''}｜"
                        f"{'阻断项' if row.get('是否阻断决策') == '是' else '非阻断项'}"
                    ),
                })
                if len(unique) == 12:
                    break
            return unique

        product_columns = {"名称", "品牌", "型号", "产品族", "系列", "核心参数"}
        if product_columns.issubset(columns):
            compact = []
            for row in rows:
                name_parts = [str(row.get(key) or "").strip() for key in ("名称", "品牌", "型号")]
                family_parts = [str(row.get(key) or "").strip() for key in ("产品族", "系列")]
                compact.append({
                    "产品 / 型号": "｜".join(dict.fromkeys(part for part in name_parts if part)),
                    "产品族 / 系列": "｜".join(dict.fromkeys(part for part in family_parts if part)),
                    "核心参数": row.get("核心参数") or "公开参数待原厂资料进一步核验",
                })
            return compact

        opportunity_columns = {"合作方向", "优先级", "切入场景", "其他公开披露事项", "立项条件"}
        if opportunity_columns.issubset(columns):
            return [{
                "合作方向 / 优先级": f"{row.get('合作方向') or ''}｜{row.get('优先级') or ''}",
                "切入场景": row.get("切入场景") or "",
                "公开依据与决策门槛": (
                    f"{row.get('其他公开披露事项') or ''}；立项条件：{row.get('立项条件') or ''}"
                ).strip("；"),
            } for row in rows]
        return rows

    @staticmethod
    def _table_widths(columns: list[str]) -> list[int]:
        """Return portrait-page widths tuned for the compact Word schemas."""
        presets = {
            ("关键事项", "重要性与影响", "建议获取资料", "时点与门槛"): [1500, 3300, 3000, 1560],
            ("产品 / 型号", "产品族 / 系列", "核心参数"): [2800, 2400, 4160],
            ("合作方向 / 优先级", "切入场景", "公开依据与决策门槛"): [2200, 2500, 4660],
            ("来源名称", "来源类型", "发布日期", "网址"): [2100, 1500, 1500, 4260],
            ("名称", "地区", "主要工艺", "运营状态"): [2500, 1700, 3460, 1700],
        }
        preset = presets.get(tuple(columns))
        if preset is not None:
            return preset
        cell_width = 9360 // max(len(columns), 1)
        return [cell_width] * len(columns)

    @staticmethod
    def _format_caption(paragraph) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.shared import Pt

        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        for run in paragraph.runs:
            run.bold = True
            run.font.size = Pt(10)

    @classmethod
    def _style_three_line_table(cls, table) -> None:
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor

        tc = theme_colors()
        navy_hex = tc["navy"].lstrip("#")
        pale_hex = "D9E2EC"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl_pr = table._tbl.tblPr
        existing = tbl_pr.find(qn("w:tblBorders"))
        if existing is not None:
            tbl_pr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge, value, size, color in (
            ("top", "single", "12", "000000"), ("bottom", "single", "12", "000000"),
            ("left", "nil", "0", "FFFFFF"), ("right", "nil", "0", "FFFFFF"),
            ("insideH", "nil", "0", "FFFFFF"), ("insideV", "nil", "0", "FFFFFF"),
        ):
            node = OxmlElement(f"w:{edge}")
            node.set(qn("w:val"), value)
            node.set(qn("w:sz"), size)
            node.set(qn("w:color"), color)
            borders.append(node)
        tbl_pr.append(borders)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                if row_index == 0:
                    shade = OxmlElement("w:shd")
                    shade.set(qn("w:fill"), pale_hex)
                    tc_pr.append(shade)
                    cell_border = OxmlElement("w:tcBorders")
                    bottom = OxmlElement("w:bottom")
                    bottom.set(qn("w:val"), "single")
                    bottom.set(qn("w:sz"), "8")
                    bottom.set(qn("w:color"), navy_hex)
                    cell_border.append(bottom)
                    tc_pr.append(cell_border)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.left_indent = Pt(0)
                    paragraph.paragraph_format.right_indent = Pt(0)
                    paragraph.paragraph_format.keep_together = True
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(word_policy()["table_size_pt"])
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
                        if row_index == 0:
                            run.bold = True
                            run.font.color.rgb = RGBColor.from_string(navy_hex)

    @staticmethod
    def _field(paragraph, instruction: str) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        run = paragraph.add_run()
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = instruction
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        text = OxmlElement("w:t")
        # Keep the initial field result empty: production refreshes it through
        # LibreOffice, while an unrefreshed draft must never expose a placeholder.
        text.text = ""
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        for node in (begin, instr, separate, text, end):
            run._r.append(node)

    @staticmethod
    def _table_geometry(table, widths: list[int]) -> None:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        table.autofit = False
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        if tbl_w is None:
            tbl_w = OxmlElement("w:tblW")
        tbl_w.set(qn("w:w"), str(sum(widths)))
        tbl_w.set(qn("w:type"), "dxa")
        if tbl_w.getparent() is None:
            tbl_pr.append(tbl_w)
        existing_layout = tbl_pr.find(qn("w:tblLayout"))
        if existing_layout is not None:
            tbl_pr.remove(existing_layout)
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_layout.set(qn("w:type"), "fixed")
        tbl_pr.append(tbl_layout)
        existing_ind = tbl_pr.find(qn("w:tblInd"))
        if existing_ind is not None:
            tbl_pr.remove(existing_ind)
        tbl_ind = OxmlElement("w:tblInd")
        tbl_ind.set(qn("w:w"), "0")
        tbl_ind.set(qn("w:type"), "dxa")
        tbl_pr.append(tbl_ind)
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
                tc_w.set(qn("w:w"), str(width))
                tc_w.set(qn("w:type"), "dxa")
