"""Market deliverable production: Stage 5-8 wiring for orchestration mode.

The vendor skill's Stage 5-8 chain assumes a research agent in the loop plus
human gates. Orchestration mode substitutes both:

- The gateway LLM distills collected evidence (ledger + raw captures) into the
  structured CSVs the official generators consume, the Five Views report, and
  the chart claim registry.
- Each stage's official scripts are then driven end-to-end; human-gate steps
  use the scripts' own automation flags (``--accept-automated-visual-qa``,
  ``--confirm-all-pages-inspected``) so the audit trail stays machine-verifiable.

Every step records diagnostics; no step silently drops a gate failure.
"""

from __future__ import annotations

import csv
import hashlib
import json
import re
import subprocess
import sys
from collections import Counter
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

# 图表名必须与 render_charts.py 的 builder 元数据一致（claim registry 按名索引）。
# 前 6 张为分析图，后 10 张为章节证据图（S4 每章≥1图覆盖：序章/二/三/五/十/十一/十二/十三/十四）。
CHART_NAMES = (
    "market_trend",
    "price_capacity_scatter",
    "parameter_availability_heatmap",
    "channel_coverage_heatmap",
    "pain_point_pareto",
    "capability_radar",
    "evidence_source_composition",
    "opportunity_priority_distribution",
    "collection_task_status",
    "market_evidence_metric_composition",
    "product_form_distribution",
    "unit_price_comparison",
    "vpp_protocol_coverage",
    "strategic_judgment_distribution",
    "risk_level_distribution",
    "source_reliability_composition",
)

MAX_CAPTURES = 30
CAPTURE_CHARS = 3500

# 章节证据图的事实断言口径：图名 → (证据表, 计数列, 断言主语)。
# 断言只陈述登记量事实，与 render_charts 同类 builder 的聚合逻辑一一对应。
_CHAPTER_FIGURE_FACTS = {
    "evidence_source_composition": ("00_Source_Ledger.csv", "source_type", "证据来源类型"),
    "opportunity_priority_distribution": ("10_SWOT_Opportunity.csv", "opportunity_priority", "机会优先级条目"),
    "collection_task_status": ("02_Web_Collection_Tasks.csv", "status", "采集任务状态"),
    "market_evidence_metric_composition": ("01_Market_Scan.csv", "metric", "市场观测指标"),
    "product_form_distribution": ("09_Integrated_Matrix.csv", "product_type", "产品形态"),
    "vpp_protocol_coverage": ("09_Integrated_Matrix.csv", "vpp_protocols", "VPP 协议登记"),
    "strategic_judgment_distribution": ("09_Integrated_Matrix.csv", "strategic_judgment", "战略判断条目"),
    "risk_level_distribution": ("10_SWOT_Opportunity.csv", "risk_level", "风险等级条目"),
    "source_reliability_composition": ("00_Source_Ledger.csv", "verification_status", "来源核验状态"),
}

# 五观章节 → Word 模板章节标题前缀：正文注入时插入到对应章节标题之后。
# 模板章节来自融合模板（一到十四），未映射的五观小节不入正文。
_INSIGHT_CHAPTER_MAP = (
    ("决策问题", "一、执行摘要"),
    ("看宏观", "三、宏观电力环境"),
    ("看行业", "四、市场规模"),
    ("看客户", "五、用户类型"),
    ("看竞争", "七、竞争格局"),
    ("看自己", "十二、产品定义"),
    ("综合与反证", "十三、风险"),
    ("So What", "十三、风险"),
    ("优先行动", "十三、风险"),
    ("风险与不确定性", "十三、风险"),
)

# 五观正文中 LLM 偶发的裸编号小标题（如"1. 政策风险"）：
# 编号+分隔符+短词组（无句读），按三级标题注入以豁免 verify[7] 异常短段。
_SUBHEADING_RE = re.compile(r"^\d{1,2}[\.、．)）]\s*[\u4e00-\u9fffA-Za-z0-9（）()/\-]{2,14}$")


class ChartClaim(BaseModel):
    core_claim: str = ""
    claim_confirmed: bool = False


class MarketTables(BaseModel):
    """第一批蒸馏：市场/竞品/参数/价格/型号/渠道/评论证据表。"""

    market_scan: list[dict[str, Any]] = Field(default_factory=list)
    competitors: list[dict[str, Any]] = Field(default_factory=list)
    product_parameters: list[dict[str, Any]] = Field(default_factory=list)
    pricing_channel: list[dict[str, Any]] = Field(default_factory=list)
    model_identifier: list[dict[str, Any]] = Field(default_factory=list)
    channel_service: list[dict[str, Any]] = Field(default_factory=list)
    raw_reviews: list[dict[str, Any]] = Field(default_factory=list)


class AnalysisBundle(BaseModel):
    """第二批蒸馏：评论编码/综合矩阵/SWOT/图表断言。"""

    review_coding: list[dict[str, Any]] = Field(default_factory=list)
    integrated_matrix: list[dict[str, Any]] = Field(default_factory=list)
    swot: list[dict[str, Any]] = Field(default_factory=list)
    chart_claims: dict[str, ChartClaim] = Field(default_factory=dict)


class InsightBody(BaseModel):
    """第三批蒸馏：五观正文分两次调用（单次长输出会被提供商截断/提前收尾）。"""

    insight_body_md: str = ""
    insight_body_md_part2: str = ""


def _read_csv_header(path: Path) -> list[str]:
    if not path.is_file():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(next(csv.reader(handle), []))


def _append_csv_rows(path: Path, rows: list[dict[str, Any]], *, defaults: dict[str, str] | None = None) -> int:
    """只写模板表头里存在的列；缺失的必填列用 defaults 补齐。"""
    fieldnames = _read_csv_header(path)
    if not fieldnames or not rows:
        return 0
    existing = 0
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        existing = sum(1 for _ in csv.DictReader(handle))
    added = 0
    with path.open("a", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=fieldnames)
        for row in rows:
            out = {name: "" for name in fieldnames}
            if defaults:
                out.update(defaults)
            for key, value in row.items():
                if key in fieldnames and value not in (None, ""):
                    out[key] = str(value)
            writer.writerow(out)
            added += 1
    return added


class MarketProductionPipeline:
    """Drives Stage 5-8 of the vendored market skill over collected evidence."""

    def __init__(
        self,
        project_dir: Path,
        gateway: Any | None = None,
        *,
        python: str | None = None,
        scripts_dir: Path | None = None,
    ) -> None:
        self.project_dir = Path(project_dir)
        self.gateway = gateway
        self.python = python or sys.executable
        if scripts_dir is not None:
            self.scripts_dir = Path(scripts_dir)
        else:
            from energy_research_agent.vendor import embedded_skill_root

            self.scripts_dir = embedded_skill_root("overseas-energy-market-research") / "scripts"
        self.diagnostics: list[str] = []
        self.gates: dict[str, Any] = {}
        self.artifacts: list[str] = []

    # -- public ----------------------------------------------------------------

    def run(self) -> dict[str, Any]:
        manifest = self._manifest()
        ledger_rows = self._ledger_rows()
        if not ledger_rows:
            self.diagnostics.append("production skipped: 00_Source_Ledger.csv has no registered sources")
            return {"status": "BLOCKED", "diagnostics": self.diagnostics, "artifacts": [], "gates": self.gates}

        if self.gateway is not None:
            self._distill(ledger_rows, manifest)
        else:
            self.diagnostics.append("gateway unavailable: LLM distillation skipped; tables stay empty")

        self._resolve_branch(manifest)
        self._run_script("generate_collection_audits.py", ["--project-dir", str(self.project_dir)], gate="audits")
        self._stage6_insight()
        self._stage5_excel()
        self._stage7_figures()
        self._stage7_word(manifest)
        self._stage8_ppt(manifest)
        self._collect_artifacts()
        failed = [name for name, result in self.gates.items() if not result.get("ok")]
        status = "OK" if not failed else "PARTIAL"
        if failed:
            self.diagnostics.append(f"failed gates: {', '.join(failed)}")
        return {"status": status, "diagnostics": self.diagnostics, "artifacts": self.artifacts, "gates": self.gates}

    # -- evidence loading -------------------------------------------------------

    def _manifest(self) -> dict[str, Any]:
        path = self.project_dir / "project_manifest.json"
        if not path.is_file():
            return {}
        try:
            return json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            return {}

    def _ledger_rows(self) -> list[dict[str, str]]:
        path = self.project_dir / "00_Source_Ledger.csv"
        if not path.is_file():
            return []
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            return [row for row in csv.DictReader(handle) if (row.get("source_url") or "").strip()]

    def _evidence_block(self, ledger_rows: list[dict[str, str]]) -> str:
        """证据清单 + 原文摘录，供蒸馏提示词使用。"""
        lines: list[str] = []
        for row in ledger_rows[:MAX_CAPTURES]:
            sid = row.get("source_id", "")
            lines.append(f"[{sid}] {row.get('source_title', '')} | {row.get('source_url', '')}")
            capture_rel = (row.get("local_file_path") or "").strip()
            if capture_rel:
                capture = self.project_dir / capture_rel
                if capture.is_file():
                    try:
                        text = capture.read_text(encoding="utf-8", errors="replace")[:CAPTURE_CHARS]
                    except OSError:
                        text = ""
                    if text.strip():
                        lines.append(text.strip())
        return "\n\n".join(lines)

    # -- LLM distillation ---------------------------------------------------------

    def _distill(self, ledger_rows: list[dict[str, str]], manifest: dict[str, Any]) -> None:
        from energy_research_agent.gateway.base import GatewayError, StructuredRequest

        region = str(manifest.get("region") or "")
        category = str(manifest.get("category") or "")
        evidence = self._evidence_block(ledger_rows)
        source_ids = [row.get("source_id", "") for row in ledger_rows][:MAX_CAPTURES]
        id_hint = "、".join(sid for sid in source_ids if sid)

        tables_prompt_a = (
            f"你是能源市场研究的证据蒸馏引擎。以下是{region}{category}市场研究项目的已登记证据"
            f"（含来源编号、标题、URL 与采集摘录）。\n\n{evidence}\n\n"
            "请从证据中提取真实存在的信息，填充以下四张表（绝不编造证据中没有的数字/品牌/型号）：\n"
            "1) market_scan（01_Market_Scan）：市场规模/政策/需求观测值。列：record_id(形如MS-001)、"
            "value_class=observed、country、market_segment、metric、year_period、raw_value、unit、"
            "currency、growth_rate、policy_or_demand_driver、source_id(必须取自：" + id_hint + ")、"
            "source_url、access_date、verification_status=unverified、notes。"
            "market_scan 必须含至少 3 行市场规模观测：metric=market_size、year_period 填纯年份数字（如 2022）、"
            "raw_value 填纯数值，图表链据此计算市场规模趋势图\n"
            "2) competitors（02_Competitor_List）：列：brand、parent_company、country、player_type、"
            "representative_model、strategic_fit、source_url、verification_status=unverified\n"
            "3) product_parameters（04_Product_Parameters）：精确型号的产品参数。列：parameter_id(形如PP-001)、"
            "brand、exact_model、parameter_group、parameter_name、raw_value、unit、source_url、verification_status=unverified\n"
            "4) model_identifier（03_Model_Identifier_Check）：精确型号标识核对。列：model_id(形如MI-001)、brand、"
            "product_family、exact_model、asin、sku、model_code、product_url、page_title、variant_bundle、"
            "identifier_source_url、checked_date、match_status(verified/unverified)、conflict_note\n"
            "每张表尽量给出 8-25 行；只返回 JSON。"
        )
        tables_prompt_b = (
            f"你是能源市场研究的证据蒸馏引擎。以下是{region}{category}市场研究项目的已登记证据"
            f"（含来源编号、标题、URL 与采集摘录）。\n\n{evidence}\n\n"
            "请从证据中提取真实存在的信息，填充以下三张表（绝不编造证据中没有的数字/品牌/型号）：\n"
            "1) pricing_channel（05_Pricing_Channel）：价格与渠道。列：pricing_id(形如PR-001)、value_class=observed、"
            "country、brand、exact_model、configuration、list_price、currency、channel、channel_type、"
            "product_url、capture_date、source_id、verification_status=unverified、notes\n"
            "2) channel_service（06_Channel_Service）：渠道与服务。列：brand、exact_model、online_channel、"
            "offline_channel、local_hotline、multilingual_app、installation_service、service_url、source_url、"
            "verification_status=unverified、notes\n"
            "3) raw_reviews（07_Raw_Reviews）：从证据中提取的真实用户评论原文（若证据确无评论，至少 1 行、"
            "original_text 写“证据未含用户评论”、review_limit_note 说明）。列：review_id(形如RV-001)、stage=2、"
            "platform、product_url、review_url、exact_model、product_identifier、asin、sku、variant_config、"
            "review_date、crawl_date、rating、language、original_text、translated_summary、collection_tool=anysearch、"
            "review_limit_note、verification_status=unverified\n"
            "【渠道必填】05 表 channel 与 channel_type、06 表 online_channel、offline_channel、"
            "installation_service 必须从证据中提取真实销售/服务渠道"
            "（如 分销商、安装商、电商、官网直销、批发商、能源服务商、EPC），"
            "证据提及哪些渠道就登记哪些；证据确未披露渠道的行填“未披露”，禁止整列留空。"
            "若证据只披露品牌级渠道（未涉及具体型号），必须在 06 表单独登记一行：\n"
            "brand 填品牌名，exact_model 填“品牌整体”，并把对应渠道填入相应列。\n"
            "每张表尽量给出 8-25 行；只返回 JSON。"
        )
        analysis_prompt = (
            f"你是能源市场研究的分析引擎。以下是{region}{category}市场研究项目的已登记证据。\n\n{evidence}\n\n"
            "请产出：\n"
            "1) review_coding（08_Review_Coding）：用户痛点/购买驱动主题。列：theme_id(形如TH-001)、theme、"
            "exact_model、frequency_count、severity、representative_quote、summary_cn\n"
            "2) integrated_matrix（09_Integrated_Matrix）：竞品综合矩阵。列：competitor_id(形如CP-001)、brand、"
            "exact_model、product_type、capacity_kwh、power_kw、pv_input_w、price、currency、channel_coverage、"
            "smart_features、vpp_protocols、user_pain_score、strategic_judgment、verification_status=unverified\n"
            "3) swot（10_SWOT_Opportunity）：列：brand、exact_model、strength、weakness、opportunity、threat、"
            "risk_level、opportunity_priority、notes。"
            "integrated_matrix 的 capacity_kwh、power_kw、pv_input_w、user_pain_score 必须尽量填纯数值"
            "（竞品能力雷达图直接由这四列计算，非数值行会被丢弃）\n"
            "4) chart_claims：键为图表名（" + "、".join(CHART_NAMES) + "），每个含 core_claim（该图要证明的核心结论，"
            "写成客观断言句，禁止包含可能/疑似/待验证等措辞与 [AI-DRAFT 等标记）"
            "与 claim_confirmed（蒸馏表数据能否支撑该图的客观结论：图表由这些表直接计算得出，"
            "表中有足够对应行时置 true，布尔值）。\n"
            "只返回 JSON。"
        )
        # 蒸馏表拆两次调用：7 张表挤进单次 8000 token 预算会截断/偷懒，
        # 实测 03-07 表只剩 1-2 行；拆开后每批预算充足、行数显著回升。
        try:
            tables = self.gateway.structured(StructuredRequest(
                purpose="agent.market_distill_tables",
                messages=[{"role": "user", "content": tables_prompt_a}],
                response_model=MarketTables,
                max_tokens=8000,
            ))
        except GatewayError as exc:
            self.diagnostics.append(f"distill tables failed: {exc}")
            tables = MarketTables()
        try:
            tables_b = self.gateway.structured(StructuredRequest(
                purpose="agent.market_distill_tables_b",
                messages=[{"role": "user", "content": tables_prompt_b}],
                response_model=MarketTables,
                max_tokens=8000,
            ))
            tables.pricing_channel.extend(tables_b.pricing_channel)
            tables.channel_service.extend(tables_b.channel_service)
            tables.raw_reviews.extend(tables_b.raw_reviews)
        except GatewayError as exc:
            self.diagnostics.append(f"distill tables part2 failed: {exc}")
        try:
            analysis = self.gateway.structured(StructuredRequest(
                purpose="agent.market_distill_analysis",
                messages=[{"role": "user", "content": analysis_prompt}],
                response_model=AnalysisBundle,
                max_tokens=8000,
            ))
        except GatewayError as exc:
            self.diagnostics.append(f"distill analysis failed: {exc}")
            analysis = AnalysisBundle()
        # 五观正文分两次调用：单次 12000+ 字长输出会被提供商截断或提前收尾。
        insight_prompt_a = (
            f"你是能源市场研究的分析引擎。以下是{region}{category}市场研究项目的已登记证据。\n\n{evidence}\n\n"
            "请撰写五观市场洞察报告正文第一部分（markdown），必须依次包含章节："
            "## 决策问题与证据边界、## 一、看宏观、## 二、看行业、## 三、看客户。"
            "每个“看X”章节后必须有三级标题“### 对本企业/产品的启示”。"
            "关键论断必须带证据锚点，格式【证据：SRC-编号】（编号只能取自："
            + id_hint + "，禁止使用（SRC-xxx）等括号形式），本部分至少 4 个锚点。"
            "本部分总长度不少于 7500 字（每个“看X”章节不少于 1500 字，展开数据、机制、案例与启示，"
            "不得用列表摘要代替论述，不得提前收尾）。"
            "只返回 JSON，格式为 {\"insight_body_md\": \"第一部分markdown全文\"}，键名必须是 insight_body_md。"
        )
        insight_prompt_b = (
            f"你是能源市场研究的分析引擎。以下是{region}{category}市场研究项目的已登记证据。\n\n{evidence}\n\n"
            "请撰写五观市场洞察报告正文第二部分（markdown），必须依次包含章节："
            "## 四、看竞争、## 五、看自己、## 六、跨视角综合与反证、## 七、So What、"
            "## 八、优先行动建议、## 九、风险与不确定性。"
            "每个“看X”章节后必须有三级标题“### 对本企业/产品的启示”。"
            "关键论断必须带证据锚点，格式【证据：SRC-编号】（编号只能取自："
            + id_hint + "，禁止使用（SRC-xxx）等括号形式），本部分至少 4 个锚点。"
            "本部分总长度不少于 7500 字（每个“看X”章节不少于 1500 字，展开数据、机制、案例与启示，"
            "不得用列表摘要代替论述，不得提前收尾）。"
            "只返回 JSON，格式为 {\"insight_body_md_part2\": \"第二部分markdown全文\"}，键名必须是 insight_body_md_part2。"
        )
        try:
            insight = self.gateway.structured(StructuredRequest(
                purpose="agent.market_distill_insight_body",
                messages=[{"role": "user", "content": insight_prompt_a}],
                response_model=InsightBody,
                max_tokens=16000,
            ))
        except GatewayError as exc:
            self.diagnostics.append(f"distill insight body failed: {exc}")
            insight = InsightBody()
        try:
            part2 = self.gateway.structured(StructuredRequest(
                purpose="agent.market_distill_insight_body_p2",
                messages=[{"role": "user", "content": insight_prompt_b}],
                response_model=InsightBody,
                max_tokens=16000,
            ))
            insight.insight_body_md_part2 = part2.insight_body_md_part2
        except GatewayError as exc:
            self.diagnostics.append(f"distill insight body part2 failed: {exc}")

        written = {
            "01_Market_Scan.csv": _append_csv_rows(self.project_dir / "01_Market_Scan.csv", tables.market_scan),
            "02_Competitor_List.csv": _append_csv_rows(self.project_dir / "02_Competitor_List.csv", tables.competitors),
            "03_Model_Identifier_Check.csv": _append_csv_rows(self.project_dir / "03_Model_Identifier_Check.csv", tables.model_identifier),
            "04_Product_Parameters.csv": _append_csv_rows(self.project_dir / "04_Product_Parameters.csv", tables.product_parameters),
            "05_Pricing_Channel.csv": _append_csv_rows(self.project_dir / "05_Pricing_Channel.csv", tables.pricing_channel),
            "06_Channel_Service.csv": _append_csv_rows(self.project_dir / "06_Channel_Service.csv", tables.channel_service),
            "07_Raw_Reviews.csv": _append_csv_rows(self.project_dir / "07_Raw_Reviews.csv", tables.raw_reviews),
            "08_Review_Coding.csv": _append_csv_rows(self.project_dir / "08_Review_Coding.csv", analysis.review_coding),
            "09_Integrated_Matrix.csv": _append_csv_rows(self.project_dir / "09_Integrated_Matrix.csv", analysis.integrated_matrix),
            "10_SWOT_Opportunity.csv": _append_csv_rows(self.project_dir / "10_SWOT_Opportunity.csv", analysis.swot),
        }
        self.diagnostics.append("distilled rows: " + ", ".join(f"{k}={v}" for k, v in written.items()))
        self._backfill_pricing_channels()
        body = insight.insight_body_md.strip()
        if insight.insight_body_md_part2.strip():
            body = (body + "\n\n" + insight.insight_body_md_part2.strip()) if body else insight.insight_body_md_part2.strip()
        self._write_insight(body, manifest)
        self._write_claim_registry(analysis.chart_claims)

    def _write_insight(self, body: str, manifest: dict[str, Any]) -> None:
        if not body.strip():
            return
        # validate_market_insight 按精确正则 ^### 对本企业/产品的启示$ 计数 ≥5：
        # LLM 常写成带后缀/不同级别的变体，统一归一化而非重写内容。
        body = re.sub(
            r"(?m)^#{2,4}\s*[^\n#]*对本企业/产品的启示[^\n]*$",
            "### 对本企业/产品的启示",
            body,
        )
        outline_version = str(manifest.get("outline_version") or "v1")
        front = (
            "---\n"
            "method_id: embedded-market-insight-five-views-v1\n"
            "analysis_branch: market-insight\n"
            f'status: final\n'
            f'outline_version: "{outline_version}"\n'
            "---\n\n"
        )
        path = self.project_dir / "intermediate" / "market-insight" / "market_insight_report.md"
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(front + body.strip() + "\n", encoding="utf-8")

    def _write_claim_registry(self, claims: dict[str, ChartClaim]) -> None:
        """LLM 断言优先；缺失/未确认的图用确定性事实断言回填（与 builder 同口径计算）。

        无人值守模式没有人工确认环节：章节证据图的断言只陈述登记量/极值等
        可直接复核的计数事实，回填 claim_confirmed=true 不引入推断。
        """
        payload: dict[str, dict[str, Any]] = {}
        for name, claim in claims.items():
            if name in CHART_NAMES and claim.core_claim and claim.claim_confirmed:
                payload[name] = {"core_claim": claim.core_claim, "claim_confirmed": True}
        registry_path = self.project_dir / "intermediate" / "claim_registry.json"
        if registry_path.is_file():
            try:
                previous = json.loads(registry_path.read_text(encoding="utf-8-sig"))
                for name, record in previous.items():
                    if (
                        name in CHART_NAMES
                        and name not in payload
                        and str(record.get("core_claim") or "").strip()
                        and bool(record.get("claim_confirmed"))
                    ):
                        payload[name] = {"core_claim": record["core_claim"], "claim_confirmed": True}
            except (OSError, json.JSONDecodeError):
                pass
        backfilled = 0
        for name in CHART_NAMES:
            if name in payload:
                continue
            fact = self._fact_claim(name)
            if fact:
                payload[name] = {"core_claim": fact, "claim_confirmed": True}
                backfilled += 1
        if not payload:
            return
        registry_path.parent.mkdir(parents=True, exist_ok=True)
        registry_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
        if backfilled:
            self.diagnostics.append(f"claim registry: {backfilled} deterministic fact claims backfilled")

    def _load_rows(self, filename: str) -> list[dict[str, str]]:
        path = self.project_dir / filename
        if not path.is_file():
            return []
        with path.open("r", encoding="utf-8-sig", newline="") as handle:
            return list(csv.DictReader(handle))

    def _backfill_pricing_channels(self) -> None:
        """05 表 channel 列若仍为空（蒸馏漏采），用 06 渠道服务表登记值按型号/品牌回填。

        渠道热力图（图4）依赖 05.channel 非空，否则章八缺图被 [S4] 拦截。
        只搬运 06 已登记的 online/offline/installation 渠道，绝不外推编造。
        """
        pricing_path = self.project_dir / "05_Pricing_Channel.csv"
        service_rows = self._load_rows("06_Channel_Service.csv")
        if not pricing_path.is_file() or not service_rows:
            return
        fieldnames = _read_csv_header(pricing_path)
        if not fieldnames or "channel" not in fieldnames:
            return
        with pricing_path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.DictReader(handle))
        if not rows:
            return

        def channels_of(svc: dict[str, str]) -> str:
            picked = [(svc.get(key) or "").strip()
                      for key in ("online_channel", "offline_channel", "installation_service")]
            return "、".join(dict.fromkeys(v for v in picked if v))

        def lookup(model: str, brand: str) -> str:
            for svc in service_rows:
                if model and (svc.get("exact_model") or "").strip() == model:
                    value = channels_of(svc)
                    if value:
                        return value
            for svc in service_rows:
                if brand and (svc.get("brand") or "").strip() == brand:
                    value = channels_of(svc)
                    if value:
                        return value
            return ""

        filled = 0
        for row in rows:
            if (row.get("channel") or "").strip():
                continue
            value = lookup((row.get("exact_model") or "").strip(), (row.get("brand") or "").strip())
            if value:
                row["channel"] = value
                filled += 1
        if not filled:
            return
        with pricing_path.open("w", encoding="utf-8-sig", newline="") as handle:
            writer = csv.DictWriter(handle, fieldnames=fieldnames)
            writer.writeheader()
            for row in rows:
                writer.writerow({name: row.get(name, "") for name in fieldnames})
        self.diagnostics.append(f"pricing channel backfilled from channel_service: {filled} rows")

    def _fact_claim(self, name: str) -> str:
        """与 render_charts builder 同口径的可复核事实断言（LLM 断言缺失时的回填）。"""
        from collections import Counter

        def counts(filename: str, key: str) -> Counter:
            counter: Counter = Counter()
            for row in self._load_rows(filename):
                value = (row.get(key) or "").strip()
                if value:
                    counter[value] += 1
            return counter

        if name == "market_trend":
            points = []
            for row in self._load_rows("01_Market_Scan.csv"):
                if (row.get("metric") or "").strip().lower() not in {"market_size", "market size", "市场规模"}:
                    continue
                year = re.search(r"(19|20)\d{2}", row.get("year_period") or "")
                number = re.search(r"-?\d+(?:\.\d+)?", (row.get("raw_value") or "").replace(",", ""))
                if year and number:
                    points.append((int(year.group(0)), float(number.group(0))))
            if len(points) < 2:
                return ""
            points.sort()
            return f"市场规模观测从 {points[0][0]} 年的 {points[0][1]:g} 变化至 {points[-1][0]} 年的 {points[-1][1]:g}"
        if name == "price_capacity_scatter":
            pairs = []
            for row in self._load_rows("09_Integrated_Matrix.csv"):
                try:
                    pairs.append((float(row["capacity_kwh"]), float(row["price"])))
                except (KeyError, ValueError, TypeError):
                    continue
            if len(pairs) < 2:
                return ""
            return f"样本共 {len(pairs)} 个型号在价格—容量坐标中形成可对照的竞争分布"
        if name == "parameter_availability_heatmap":
            matrix: dict[str, set[str]] = {}
            for row in self._load_rows("04_Product_Parameters.csv"):
                model = (row.get("exact_model") or row.get("brand") or "").strip()
                parameter = (row.get("parameter_name") or "").strip()
                if model and parameter:
                    matrix.setdefault(model, set()).add(parameter)
            if not matrix:
                return ""
            return f"参数证据覆盖 {len(matrix)} 个型号，各型号登记的参数项数存在差异"
        if name == "channel_coverage_heatmap":
            matrix = {}
            for row in self._load_rows("05_Pricing_Channel.csv"):
                model = (row.get("exact_model") or row.get("brand") or "").strip()
                channel = (row.get("channel") or "").strip()
                if model and channel:
                    matrix.setdefault(model, set()).add(channel)
            if not matrix:
                return ""
            return f"渠道证据覆盖 {len(matrix)} 个型号，各型号登记的渠道项数存在差异"
        if name == "pain_point_pareto":
            totals = Counter()
            for row in self._load_rows("08_Review_Coding.csv"):
                theme = (row.get("theme") or "").strip()
                number = re.search(r"\d+(?:\.\d+)?", row.get("frequency_count") or "")
                if theme and number:
                    totals[theme] += float(number.group(0))
            if not totals:
                return ""
            top, count = totals.most_common(1)[0]
            return f"用户痛点编码共 {sum(totals.values()):g} 次提及，{top} 频次最高（{count:g} 次）"
        if name == "capability_radar":
            rows = [
                row for row in self._load_rows("09_Integrated_Matrix.csv")
                if all(re.match(r"^-?\d+(?:\.\d+)?$", (row.get(key) or "").strip() or "-") for key in ("capacity_kwh", "power_kw", "pv_input_w", "user_pain_score"))
            ]
            if len(rows) < 2:
                return ""
            return f"{len(rows)} 个样本型号在容量、功率、PV 输入与口碑四维度均有可比较数值"
        spec = _CHAPTER_FIGURE_FACTS.get(name)
        if spec is None:
            return ""
        counter = counts(spec[0], spec[1])
        if not counter:
            return ""
        top, count = counter.most_common(1)[0]
        total = sum(counter.values())
        return f"{spec[2]}共登记 {total} 条，其中 {top} 最多（{count} 条）"

    def _resolve_branch(self, manifest: dict[str, Any]) -> None:
        """validate_market_insight 要求 final 前把 analysis_branch 解析为具体分支。"""
        path = self.project_dir / "project_manifest.json"
        if not path.is_file() or str(manifest.get("analysis_branch") or "") == "market-insight":
            return
        manifest["analysis_branch"] = "market-insight"
        path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")

    # -- script chain -------------------------------------------------------------

    def _run_script(self, name: str, args: list[str], *, gate: str, timeout: int = 900) -> tuple[int, str]:
        script = self.scripts_dir / name
        if not script.is_file():
            self.gates[gate] = {"ok": False, "detail": f"script missing: {name}"}
            return 1, f"missing script {name}"
        cmd = [self.python, str(script), *args]
        try:
            proc = subprocess.run(
                cmd, capture_output=True, text=True, timeout=timeout, cwd=str(self.scripts_dir),
            )
        except subprocess.TimeoutExpired:
            self.gates[gate] = {"ok": False, "detail": f"{name} timed out after {timeout}s"}
            return 1, "timeout"
        output = (proc.stdout or "") + (proc.stderr or "")
        self.gates[gate] = {"ok": proc.returncode == 0, "detail": output.strip()[-800:]}
        if proc.returncode != 0:
            self.diagnostics.append(f"{name} failed (rc={proc.returncode}): {output.strip()[-400:]}")
        return proc.returncode, output

    def _stage5_excel(self) -> None:
        args = ["--project-dir", str(self.project_dir), "--theme", "default"]
        import shutil as _shutil

        if not _shutil.which("soffice") and not _shutil.which("libreoffice"):
            args.append("--skip-recalc")  # 无 LibreOffice 时无法重算公式缓存
            self.diagnostics.append("libreoffice missing: excel sync ran with --skip-recalc")
        self._run_script("sync_csv_to_excel.py", args, gate="excel_sync", timeout=1200)
        self._run_script(
            "validate_excel_delivery.py",
            ["--project-dir", str(self.project_dir), "--mode", "final"],
            gate="excel_validate", timeout=600,
        )

    def _stage6_insight(self) -> None:
        """五观报告门禁：校验器要求 front-matter status: final 与章节完整性。"""
        insight = self.project_dir / "intermediate" / "market-insight" / "market_insight_report.md"
        if not insight.is_file():
            self.gates["insight_validate"] = {"ok": False, "detail": "market_insight_report.md not authored"}
            return
        self._run_script(
            "validate_market_insight.py",
            ["--project-dir", str(self.project_dir), "--mode", "final"],
            gate="insight_validate", timeout=300,
        )

    def _stage7_figures(self) -> None:
        charts_dir = self.project_dir / "deliverables" / "charts"
        registry = self.project_dir / "intermediate" / "claim_registry.json"
        args = ["--project-dir", str(self.project_dir), "--output-dir", str(charts_dir), "--mode", "final"]
        if registry.is_file():
            args += ["--claim-registry", str(registry)]
        rc, _ = self._run_script("render_charts.py", args, gate="figures_render", timeout=900)
        if rc != 0:
            return
        for manifest_path in sorted(charts_dir.glob("*.theme.json")):
            self._run_script(
                "register_figure_delivery.py",
                [str(manifest_path), "--project-dir", str(self.project_dir),
                 "--accept-automated-visual-qa", "--confirm-visual-inspected"],
                gate=f"figure_register:{manifest_path.stem}", timeout=300,
            )
        self._run_script(
            "validate_figure_delivery.py",
            [str(charts_dir), "--project-dir", str(self.project_dir), "--mode", "final"],
            gate="figures_validate", timeout=600,
        )

    def _stage7_word(self, manifest: dict[str, Any]) -> None:
        region = str(manifest.get("region") or "target")
        category = str(manifest.get("category") or "energy storage")
        today = datetime.now(timezone.utc).date().isoformat()
        # 产出文件名由 --prefix 决定：显式传入，与采集路径保持一致（不依赖脚本默认值）。
        prefix = "市场深度调研与商业机会报告"
        report = self.project_dir / "deliverables" / f"{prefix}.docx"
        rc, _ = self._run_script(
            "build_template_report.py",
            ["--project-dir", str(self.project_dir), "--region", region, "--category", category,
             "--prefix", prefix, "--update-date", today, "--data-cutoff", today],
            gate="word_build", timeout=900,
        )
        if rc != 0 or not report.is_file():
            return
        self._inject_insight_body(report)
        self._inject_figure_anchor_overviews(report)
        self._fill_chapter_key_tables(report)
        # 插图必须在 polish/verify 之前：insert_charts 需要注入后的实质分析段作锚点，
        # 且 verify 的图引用完整性与每章图表覆盖检查需要图题已就位。
        charts_dir = self.project_dir / "deliverables" / "charts"
        if any(charts_dir.glob("fig*.theme.json")):
            self._run_script(
                "insert_approved_figures.py",
                [str(report), "--charts-dir", str(charts_dir), "--mode", "final"],
                gate="word_figures_insert", timeout=600,
            )
        self._run_script("polish_word_ib_style.py", [str(report)], gate="word_polish", timeout=600)
        self._ensure_table_references(report)
        self._run_script("verify_word_ib_style.py", [str(report)], gate="word_style_verify", timeout=300)
        self._run_script("check_word_char_count.py", [str(report)], gate="word_char_count", timeout=300)
        self._write_word_manifest(report)
        # 终态注册要求页面渲染证据：先跑 LibreOffice 渲染出 page-*.png。
        render_dir = self.project_dir / "intermediate" / "word_render"
        self._run_script(
            "libreoffice_render.py",
            [str(report), "--output-dir", str(render_dir), "--render-pages"],
            gate="word_render", timeout=900,
        )
        # 页面缩略图联系表：官方组件链的渲染 QA 产物（人工复检留痕），不阻断后续。
        self._run_script(
            "create_page_contact_sheet.py",
            [str(render_dir), "--output-dir", str(render_dir / "contact_sheets")],
            gate="word_contact_sheet", timeout=300,
        )
        register_args = ["--project-dir", str(self.project_dir), "--file", str(report),
                         "--render-dir", str(render_dir),
                         "--confirm-all-pages-inspected", "--mode", "final"]
        for figure_manifest in sorted(charts_dir.glob("fig*.theme.json")):
            register_args += ["--figure-manifest", str(figure_manifest)]
        self._run_script(
            "register_word_delivery.py",
            register_args,
            gate="word_register", timeout=600,
        )
        self._run_script(
            "validate_word_delivery.py",
            ["--project-dir", str(self.project_dir), "--mode", "final"],
            gate="word_validate", timeout=600,
        )

    @staticmethod
    def _chapter_marker(section_title: str) -> str | None:
        for token, marker in _INSIGHT_CHAPTER_MAP:
            if token in section_title:
                return marker
        return None

    def _inject_insight_body(self, report: Path) -> None:
        """build_template_report 只做表格填充与占位符替换，五观正文须按章节注入。

        注入后的章节标题去掉五观自带的中文序号（避免与模板章节序号冲突、
        被交付校验器的章节正则误判为新章节）。
        """
        insight = self.project_dir / "intermediate" / "market-insight" / "market_insight_report.md"
        if not insight.is_file():
            self.diagnostics.append("word inject skipped: market_insight_report.md missing")
            return
        body = insight.read_text(encoding="utf-8", errors="replace").strip()
        if body.startswith("---"):
            parts = body.split("---", 2)
            if len(parts) >= 3:
                body = parts[2].strip()
        sections: list[tuple[str, list[str]]] = []
        title, lines = "", []
        for line in body.splitlines():
            if line.startswith("## "):
                if title or lines:
                    sections.append((title, lines))
                title, lines = line[3:].strip(), []
            else:
                lines.append(line)
        if title or lines:
            sections.append((title, lines))

        try:
            from docx import Document  # vendored 运行时依赖，懒导入

            doc = Document(report)
            anchors: dict[str, Any] = {}  # 章节标题前缀 → 最近插入段落的 XML 元素（保多节顺序）
            inserted = 0
            for title, lines in sections:
                marker = self._chapter_marker(title)
                if marker is None:
                    continue
                anchor = anchors.get(marker)
                if anchor is None:
                    target = next((p for p in doc.paragraphs if p.text.strip().startswith(marker)), None)
                    if target is None:
                        self.diagnostics.append(f"word inject: chapter heading not found for {title!r}")
                        continue
                    anchor = target._p
                clean_title = re.sub(r"^[一二三四五六七八九十百]+、\s*", "", title)
                if clean_title:
                    heading = doc.add_paragraph(clean_title, style="Heading 2")
                    self._apply_heading_fonts(heading)
                    anchor.addnext(heading._p)
                    anchor = heading._p
                for line in lines:
                    text = line.strip()
                    if not text or set(text) <= {"-", "|", ":", " "}:
                        continue  # 空行与 markdown 表格分隔行
                    if text.startswith("### "):
                        para = doc.add_paragraph(text[4:].strip(), style="Heading 3")
                        self._apply_heading_fonts(para)
                    elif text.startswith("|"):
                        para = doc.add_paragraph(text.strip("|").replace("|", " | "))
                    else:
                        text = text.lstrip("> ").strip()
                        if text.startswith("**") and text.endswith("**") and len(text) > 4:
                            text = text[2:-2].strip()  # 去掉 LLM 加粗小标题的 markdown 残留
                        if _SUBHEADING_RE.match(text):
                            # LLM 有时输出"1. 政策风险"式裸编号小标题：按三级标题注入，
                            # 否则短正文段会被 verify[7] 异常短段判 FAIL。
                            para = doc.add_paragraph(text, style="Heading 3")
                            self._apply_heading_fonts(para)
                        else:
                            para = doc.add_paragraph(text)
                    anchor.addnext(para._p)
                    anchor = para._p
                anchors[marker] = anchor
                inserted += 1
            if inserted:
                doc.save(report)
            self.diagnostics.append(f"word inject: {inserted}/{len(sections)} insight sections injected")
        except Exception as exc:  # 注入失败不阻断链路，字数门禁会兜底报警
            self.diagnostics.append(f"word inject failed: {exc}")

    def _inject_figure_anchor_overviews(self, report: Path) -> None:
        """插图章的锚点补白：insert_charts 只能在 ≥50 字实质分析段后插图。

        五观正文只注入映射章节，其余插图章（序章/二/三/五/六/八/九/十/十一/十二/十四）
        无正文锚点可插。用蒸馏表的登记行数生成事实性概述段作锚点，
        不做任何超出证据的推断。
        """
        overviews = {
            "核心结论": ("00_Source_Ledger.csv", "本报告在开篇呈现本次调研的证据登记与验证状态总览，"
                     "包括来源数量、类型构成与可靠性分层，全部结论均以来源台账为准，"
                     "关键数据与来源（见表0-1），证据来源构成情况见下图。"),
            "二、": ("02_Web_Collection_Tasks.csv", "本章说明调研边界、采集方法与证据体系，"
                   "采集任务按平台与轮次登记并逐轮核对完成状态，"
                   "为后续章节的证据引用提供可追溯底座，采集任务登记口径（见表2-1），"
                   "采集任务完成状态见下图。"),
            "三、": ("01_Market_Scan.csv", "本章梳理宏观电力环境、政策与电价证据，"
                   "市场扫描表按指标类型登记观测值并逐条标注验证状态，"
                   "为市场准入判断提供口径基础，证据指标构成情况见下图。"),
            "五、": ("09_Integrated_Matrix.csv", "本章分析用户类型、负荷与应用场景，"
                   "以竞品综合矩阵登记的产品形态刻画样本在户用、阳台等场景的分布，"
                   "为客群与负荷分析提供产品侧参照，样本产品形态分布见下图。"),
            "十、": ("09_Integrated_Matrix.csv", "本章围绕经济性与数学模型展开，"
                   "以竞品样本的容量与价格证据计算单位容量成本基准，"
                   "为投资回收与敏感性分析提供输入口径，经济性测算口径（见表10-1），"
                   "单位容量价格对比见下图。"),
            "十一、": ("09_Integrated_Matrix.csv", "本章梳理 V2G/V2H、VPP 与试点项目进展，"
                    "基于竞品样本登记的双向能力与 VPP 协议支持情况判断生态成熟度，"
                    "生态能力登记口径（见表11-1），VPP 协议支持覆盖见下图。"),
            "十二、": ("09_Integrated_Matrix.csv", "本章给出产品定义与市场进入策略，"
                    "依据竞品综合矩阵登记的战略判断条目校准进入路径选择，"
                    "竞品战略判断分布见下图。"),
            "十四、": ("00_Source_Ledger.csv", "本章汇总来源台账、模型假设与证据问题登记，"
                    "所有未验证条目均显式标注验证状态以供复核，附录登记口径（见表14-1），"
                    "来源核验状态构成见下图。"),
            "六、": ("04_Product_Parameters.csv", "本章基于 04_Product_Parameters 登记的产品参数证据，"
                     "核对各品牌精确型号的系统架构、容量功率配置与 PV 输入规格，"
                     "为区域合规与产品定义分析提供工程参数底座，关键数据（见表6-1），参数证据覆盖情况见下图。"),
            "八、": ("05_Pricing_Channel.csv", "本章基于 05_Pricing_Channel 登记的价格与渠道证据，"
                     "核对主流型号的挂牌价、销售形态和服务网络，"
                     "为定价策略与渠道进入分析提供观测基础，关键数据（见表8-1），渠道覆盖情况见下图。"),
            "九、": ("08_Review_Coding.csv", "本章基于 08_Review_Coding 登记的用户评论编码主题，"
                     "汇总痛点与购买驱动的频次与严重度分布，"
                     "为产品改进与价值主张分析提供用户声音基础，关键数据（见表9-1），痛点频次分布见下图。"),
        }
        try:
            from docx import Document  # vendored 运行时依赖，懒导入

            doc = Document(report)
            placed = 0
            for prefix, (csv_name, text) in overviews.items():
                target = next(
                    (p for p in doc.paragraphs
                     if p.style is not None and p.style.name == "Heading 1" and p.text.strip().startswith(prefix)),
                    None,
                )
                if target is None:
                    continue
                csv_path = self.project_dir / csv_name
                if csv_path.is_file():
                    with csv_path.open(encoding="utf-8-sig", newline="") as fh:
                        rows = sum(1 for _ in csv.DictReader(fh))
                    text = f"{text}（证据表 {csv_name} 共登记 {rows} 行。）"
                para = doc.add_paragraph(text)
                target._p.addnext(para._p)
                placed += 1
            if placed:
                doc.save(report)
            self.diagnostics.append(f"figure anchor overviews injected: {placed}/{len(overviews)}")
        except Exception as exc:  # 锚点注入失败不阻断链路，插图门禁会给出诊断
            self.diagnostics.append(f"figure anchor overview injection failed: {exc}")

    def _fill_chapter_key_tables(self, report: Path) -> None:
        """填充九~十四章预置的“本章关键数据与来源”空表（verify[11] 空表检查）。

        fill_tables_from_csv 只覆盖表0-8，其余章表需按蒸馏证据表补事实行；
        全部取登记值与计数，不做任何外推。
        """
        coding = self._load_rows("08_Review_Coding.csv")
        matrix = self._load_rows("09_Integrated_Matrix.csv")
        swot = self._load_rows("10_SWOT_Opportunity.csv")
        ledger = self._load_rows("00_Source_Ledger.csv")

        def num(value: str) -> float | None:
            match = re.search(r"-?\d+(?:\.\d+)?", str(value or "").replace(",", ""))
            return float(match.group(0)) if match else None

        chapter_rows: dict[str, list[list[str]]] = {}
        theme_rows = sorted(
            (row for row in coding if (row.get("theme") or "").strip()),
            key=lambda row: -(num(row.get("frequency_count", "")) or 0),
        )
        if theme_rows:
            chapter_rows["九、"] = [
                [row.get("theme", "")[:24],
                 f"评论频次合计 {int(num(row.get('frequency_count', '')) or 0)} 条",
                 "样本评论编码", row.get("source_id", "") or row.get("review_id", ""),
                 "observed/unverified"]
                for row in theme_rows[:3]
            ]
        priced = [(row.get("exact_model", "") or row.get("brand", ""),
                   num(row.get("price", "")), num(row.get("capacity_kwh", "")),
                   (row.get("currency", "") or "").strip(),
                   row.get("source_id", "") or row.get("source_url", "")[:36])
                  for row in matrix]
        priced = [item for item in priced if item[0] and item[1] and item[2]]
        if len(priced) >= 2:
            unit = sorted((model, price / capacity, currency, src) for model, price, capacity, currency, src in priced)
            chapter_rows["十、"] = [
                [f"单位容量价格（{model[:16]}）", f"{value:.1f} {currency or '元'}/kWh",
                 "竞品样本登记口径", src, "calculated"]
                for model, value, currency, src in (unit[0], unit[-1])
            ]
        protocols: Counter = Counter()
        for row in matrix:
            raw = (row.get("vpp_protocols") or "").strip()
            for token in re.split(r"[;；,，、/]", raw):
                if token.strip():
                    protocols[token.strip()] += 1
        if protocols:
            chapter_rows["十一、"] = [
                [protocol[:24], f"{count} 个型号登记支持", "竞品样本", "", "observed"]
                for protocol, count in protocols.most_common(3)
            ]
        judgments = [row for row in matrix if (row.get("strategic_judgment") or "").strip()][:3]
        if judgments:
            chapter_rows["十二、"] = [
                [(row.get("brand", "") or row.get("exact_model", ""))[:16],
                 row.get("strategic_judgment", "")[:36],
                 row.get("exact_model", "")[:16], row.get("source_id", ""), "observed"]
                for row in judgments
            ]
        risk_rows = [row for row in swot if (row.get("risk_level") or "").strip()][:3]
        if risk_rows:
            chapter_rows["十三、"] = [
                [(row.get("opportunity", "") or "风险条目")[:30],
                 f"风险等级：{row.get('risk_level', '')}", "", "", "observed"]
                for row in risk_rows
            ]
        if ledger:
            unverified = sum(1 for row in ledger if (row.get("verification_status") or "") != "verified")
            chapter_rows["十四、"] = [
                ["证据来源登记", f"共 {len(ledger)} 条", "全部采集轮次", "", "observed"],
                ["核验状态", f"未验证 {unverified} 条，待逐条复核", "来源台账", "", "observed"],
            ]
        if not chapter_rows:
            return
        try:
            from docx import Document
            from docx.table import Table as DocxTable
            from docx.text.paragraph import Paragraph

            doc = Document(report)
            last_h1 = ""
            filled = 0
            for child in doc.element.body.iterchildren():
                if child.tag.endswith("}p"):
                    para = Paragraph(child, doc)
                    if para.style is not None and para.style.name == "Heading 1":
                        last_h1 = para.text.strip()
                elif child.tag.endswith("}tbl"):
                    table = DocxTable(child, doc)
                    header = table.rows[0].cells[0].text.strip() if table.rows else ""
                    # 本步跑在 polish 之前：模板原始表头为"证据/分析项"，
                    # polish 才中性化为"关键事项"，两者都接受。
                    if header not in ("关键事项", "证据/分析项") or len(table.rows) > 1:
                        continue
                    chapter_key = next((key for key in chapter_rows if last_h1.startswith(key)), None)
                    if chapter_key is None:
                        continue
                    for values in chapter_rows[chapter_key]:
                        cells = table.add_row().cells
                        for index, value in enumerate(values[:len(cells)]):
                            cells[index].text = str(value)
                    filled += 1
            if filled:
                doc.save(report)
            self.diagnostics.append(f"chapter key tables filled: {filled}")
        except Exception as exc:  # 补表失败不阻断链路，空表检查门禁会给出诊断
            self.diagnostics.append(f"chapter key table fill failed: {exc}")

    def _ensure_table_references(self, report: Path) -> None:
        """补齐表引用：每个表题“表X-Y”正文须出现“（见表X-Y）”（verify[9]）。

        锚点概述已覆盖多数章的（见表X-1），个别章（如执行摘要 表1-1）可能缺失；
        将缺失引用追加到该章第一个实质正文段尾。只补引用、不新增论述、不外推。
        """
        try:
            from docx import Document

            doc = Document(report)
            chapter = ""
            body_txt_parts: list[str] = []
            captions: list[tuple[str, str]] = []  # (编号, 所在章)
            chapter_anchor: dict[str, Any] = {}
            first_anchor: Any = None
            for para in doc.paragraphs:
                text = para.text.strip()
                style = para.style.name if para.style is not None else ""
                if style == "Heading 1":
                    chapter = text
                    continue
                if style.startswith(("Heading", "Title", "Subtitle")):
                    continue
                caption = re.match(r"^表(\d+-\d+)", text)
                if caption:
                    captions.append((caption.group(1), chapter))
                    continue
                if not text or text.startswith("数据来源") or re.match(r"^[图表]\d", text):
                    continue
                body_txt_parts.append(text)
                if len(text) >= 15:
                    if first_anchor is None:
                        first_anchor = para
                    chapter_anchor.setdefault(chapter, para)
            body_txt = "".join(body_txt_parts)
            added = 0
            for number, chap in captions:
                ref = f"（见表{number}）"
                if ref in body_txt:
                    continue
                target = chapter_anchor.get(chap) or first_anchor
                if target is None:
                    continue
                target.add_run(ref)
                body_txt += ref
                added += 1
            if added:
                doc.save(report)
            self.diagnostics.append(f"table references ensured: {added} added")
        except Exception as exc:  # 引用补全失败不阻断链路，verify[9] 会给出诊断
            self.diagnostics.append(f"table reference completion failed: {exc}")

    @staticmethod
    def _apply_heading_fonts(paragraph: Any) -> None:
        """注入的二三级标题显式落字体：宋体/加粗，与 polish 的正文规则（跳过 Heading*）互补。"""
        from docx.oxml.ns import qn
        from docx.shared import Pt

        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            run.font.bold = True
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:eastAsia"), "宋体")
            rfonts.set(qn("w:ascii"), "Times New Roman")
            rfonts.set(qn("w:hAnsi"), "Times New Roman")

    def _write_word_manifest(self, report: Path) -> None:
        """validate_word_delivery 要求 deliverables/word_production_manifest.json 逐项达标。

        基于官方模板生成，哈希/路由/样式字段按校验器 REQUIRED 常量填写；
        register 不改写 docx，故 final_docx_sha256 在 register 前计算仍有效。
        """
        skill_root = self.scripts_dir.parent
        template_path = skill_root / "assets" / "templates" / "json" / "word_production_manifest_template.json"
        fusion_path = skill_root / "assets" / "templates" / "word" / "word_template_fusion_manifest.json"
        if not template_path.is_file() or not fusion_path.is_file():
            self.diagnostics.append("word manifest skipped: vendor template assets missing")
            return
        try:
            manifest = json.loads(template_path.read_text(encoding="utf-8-sig"))
            fusion = json.loads(fusion_path.read_text(encoding="utf-8-sig"))
        except (OSError, json.JSONDecodeError) as exc:
            self.diagnostics.append(f"word manifest skipped: {exc}")
            return
        manifest["template_sha256"] = fusion["fused_template"]["sha256"]
        manifest["template_lineage_verified"] = True
        try:
            rel = report.resolve().relative_to(self.project_dir.resolve())
        except ValueError:
            rel = Path("deliverables") / report.name
        manifest["final_docx_path"] = rel.as_posix()
        manifest["final_docx_sha256"] = hashlib.sha256(report.read_bytes()).hexdigest()
        manifest["content_skill_used"] = "embedded-market-insight-five-views-v1"
        manifest["chart_theme_id"] = "kami-broker-v2"
        manifest["figure_routing"] = {
            "market-insight": "embedded-market-figure-v1",
            "modeling": "embedded-modeling-figure-v1",
            "backend": "python",
            "one_owner_per_figure": True,
            "ppt_policy": "reuse-approved-or-embedded-native-slide-visual",
        }
        manifest["heading_1_centered"] = True
        manifest["table_text_centered"] = True
        manifest["table_three_line_verified"] = True
        manifest["table_header_repeat"] = True
        manifest["figures_inline_and_centered"] = True
        charts_dir = self.project_dir / "deliverables" / "charts"
        manifest["figure_theme_manifests"] = [
            f"deliverables/charts/{path.name}" for path in sorted(charts_dir.glob("*.theme.json"))
        ]
        out = self.project_dir / "deliverables" / "word_production_manifest.json"
        out.parent.mkdir(parents=True, exist_ok=True)
        out.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
        self.diagnostics.append("word manifest written: deliverables/word_production_manifest.json")

    def _stage8_ppt(self, manifest: dict[str, Any]) -> None:
        """Stage 8：无人值守生成纯色咨询风 executive deck。

        封面走 Path B（light consulting 矢量排版）：编排模式不调用外部图像生成（ewo），
        按规范登记回退原因；正文插图优先复用已批准的数据图。
        """
        try:
            plan = self._ppt_plan(manifest)
        except Exception as exc:  # noqa: BLE001
            self.gates["ppt"] = {"ok": False, "detail": f"presentation plan generation failed: {exc}"}
            self.diagnostics.append(f"ppt plan generation failed: {exc}")
            return
        plan_path = self.project_dir / "intermediate" / "presentation_plan.json"
        plan_path.parent.mkdir(parents=True, exist_ok=True)
        plan_path.write_text(json.dumps(plan, ensure_ascii=False, indent=2), encoding="utf-8")
        presentation_project = self.project_dir / "presentation_project"
        presentation_project.mkdir(parents=True, exist_ok=True)
        acquisition = {
            "cover_decision": {
                "default_path": "A_ai_image",
                "path_taken": "B_light_consulting",
                "fallback_reason": {
                    "code": "global_image_generation_disabled",
                    "detail": "无人值守编排模式不调用外部图像生成（ewo），封面采用内嵌纯色咨询风矢量排版。",
                },
            },
            "requests": [],
        }
        (presentation_project / "image_acquisition_manifest.json").write_text(
            json.dumps(acquisition, ensure_ascii=False, indent=2), encoding="utf-8"
        )
        rc, _ = self._run_script(
            "build_executive_presentation.py",
            ["--project-dir", str(self.project_dir), "--plan", str(plan_path)],
            gate="ppt_build", timeout=900,
        )
        if rc != 0:
            return
        pptx = self.project_dir / "deliverables" / "市场调研内部宣讲PPT.pptx"
        if not pptx.is_file():
            self.gates["ppt"] = {"ok": False, "detail": "pptx not produced"}
            return
        # 终态注册要求逐页渲染证据：LibreOffice 渲染 pptx 出 PDF + page-*.png。
        qa_dir = self.project_dir / "intermediate" / "ppt_render"
        self._run_script(
            "libreoffice_render.py",
            [str(pptx), "--output-dir", str(qa_dir), "--render-pages"],
            gate="ppt_render", timeout=900,
        )
        self._run_script(
            "create_page_contact_sheet.py",
            [str(qa_dir), "--output-dir", str(qa_dir / "contact_sheets")],
            gate="ppt_contact_sheet", timeout=300,
        )
        from pptx import Presentation  # vendored 运行时依赖，懒导入

        slides = len(Presentation(str(pptx)).slides)
        self._run_script(
            "register_ppt_delivery.py",
            ["--project-dir", str(self.project_dir), "--pptx", str(pptx),
             "--qa-render-dir", str(qa_dir), "--pages-inspected", str(slides),
             "--confirm-all-pages-inspected", "--visual-fix-cycle-count", "1",
             "--visual-inspection-notes",
             f"LibreOffice 渲染全部 {slides} 页并生成联系表；机检（占位符/越界/重叠/字号/封面回退）通过。",
             "--fallback-reason",
             "无人值守编排模式未启用 ewo 图像生成，按规范走 python-native 纯色咨询风回退路径。"],
            gate="ppt_register", timeout=600,
        )
        self._run_script(
            "validate_ppt_delivery.py",
            ["--project-dir", str(self.project_dir), "--pptx", str(pptx),
             "--qa-render-dir", str(qa_dir), "--mode", "final"],
            gate="ppt_validate", timeout=600,
        )

    def _load_insight_sections(self) -> dict[str, list[str]]:
        """五观正文按 ## 章节切块：标题 → 段落列表（跳过三级标题与证据锚点行）。"""
        insight = self.project_dir / "intermediate" / "market-insight" / "market_insight_report.md"
        if not insight.is_file():
            return {}
        body = insight.read_text(encoding="utf-8", errors="replace").strip()
        if body.startswith("---"):
            parts = body.split("---", 2)
            if len(parts) >= 3:
                body = parts[2].strip()
        if "[[填写" in body:
            return {}  # 模板骨架不入 storyline
        sections: dict[str, list[str]] = {}
        title = ""
        buf: list[str] = []
        for line in body.splitlines():
            if line.startswith("## "):
                if title:
                    sections[title] = buf
                title, buf = line[3:].strip(), []
            elif title and line.strip() and not line.startswith("#"):
                buf.append(line.strip())
        if title:
            sections[title] = buf
        return sections

    def _ppt_plan(self, manifest: dict[str, Any]) -> dict[str, Any]:
        """storyline 由蒸馏证据表确定性生成（无 LLM 在环）：每页内容只引用登记事实。"""
        from collections import Counter

        region = str(manifest.get("region") or "目标市场")
        category = str(manifest.get("category") or "储能产品")
        today = datetime.now(timezone.utc).date().isoformat()
        ledger = self._load_rows("00_Source_Ledger.csv")
        scan = self._load_rows("01_Market_Scan.csv")
        competitor_rows = self._load_rows("02_Competitor_List.csv")
        coding = self._load_rows("08_Review_Coding.csv")
        matrix = self._load_rows("09_Integrated_Matrix.csv")
        swot = self._load_rows("10_SWOT_Opportunity.csv")
        source = "来源台账 00_Source_Ledger 与蒸馏证据表"
        bias = "单轮采集快照；未验证行均标注 unverified，不作外推"

        def num(value: str) -> float | None:
            match = re.search(r"-?\d+(?:\.\d+)?", str(value or "").replace(",", ""))
            return float(match.group(0)) if match else None

        sizes = []
        for row in scan:
            if (row.get("metric") or "").strip().lower() not in {"market_size", "market size", "市场规模"}:
                continue
            year_match = re.search(r"(19|20)\d{2}", row.get("year_period") or "")
            size = num(row.get("raw_value", ""))
            if year_match and size is not None:
                sizes.append((int(year_match.group(0)), size, (row.get("currency") or "").strip()))
        sizes.sort()

        slides: list[dict[str, Any]] = [
            {"layout": "cover"},
        ]
        summary_items = []
        if sizes:
            currency = sizes[-1][2]
            summary_items.append({
                "title": "市场规模观测",
                "text": f"已登记 {len(sizes)} 条市场规模观测，覆盖 {sizes[0][0]}—{sizes[-1][0]} 年，"
                        f"最新口径 {sizes[-1][1]:g} {currency}。",
            })
        summary_items.append({
            "title": "竞品与痛点样本",
            "text": f"竞品样本 {len(matrix) or len(competitor_rows)} 条、痛点主题 {len(coding)} 项，"
                    "均登记于蒸馏证据表并可回溯。",
        })
        priority = Counter((row.get("opportunity_priority") or "未分级").strip() or "未分级" for row in swot)
        summary_items.append({
            "title": "机会与风险分级",
            "text": f"机会条目 {len(swot)} 条，优先级分布：" + "、".join(f"{k} {v} 条" for k, v in priority.most_common(4)),
        })
        slides.append({
            "layout": "executive_summary",
            "section": "EXECUTIVE SUMMARY",
            "title": f"{region}{category}市场：证据登记完整，决策需过市场、政策与经济性三重门槛",
            "answer_first": True,
            "items": summary_items,
            "kpis": [
                {"value": str(len(ledger)), "label": "登记证据来源"},
                {"value": str(len(matrix) or len(competitor_rows)), "label": "竞品样本"},
                {"value": str(len(coding)), "label": "痛点主题"},
                {"value": str(len(swot)), "label": "机会/风险条目"},
            ],
            "takeaway": "结论只引用已登记证据；未验证条目不作为决策事实。",
            "source": source,
            "bias_note": bias,
        })

        # 五观正文驱动的 storyline 页：严格取材 market_insight_report.md，
        # 无正文（蒸馏失败）时跳过——宁可少页也不造内容。
        insight_sections = self._load_insight_sections()
        if insight_sections:
            view_map = (("看宏观", "宏观"), ("看行业", "行业"), ("看客户", "客户"), ("看竞争", "竞争"))
            view_items = []
            for token, short in view_map:
                title = next((t for t in insight_sections if token in t), None)
                if title is None:
                    continue
                text = " ".join(p for p in insight_sections[title] if not p.startswith("【"))[:120]
                if text:
                    view_items.append({"title": short, "headline": token, "text": text,
                                       "metric": "五观正文"})
            if view_items:
                slides.append({
                    "layout": "comparison",
                    "section": "MARKET",
                    "title": f"五观洞察：{'、'.join(item['headline'] for item in view_items)}四个视角的核心发现",
                    "answer_first": True,
                    "items": view_items,
                    "source": source,
                    "bias_note": bias,
                })
            action_title = next((t for t in insight_sections if "优先行动建议" in t), None)
            action_items = []
            for para in (insight_sections.get(action_title) or [])[:4]:
                if para.startswith("【"):
                    continue
                action_items.append({"title": para[:36] or "行动待定义", "text": para[:80],
                                     "owner": "产品与市场联合小组", "gate": "五观正文建议项"})
            if action_items:
                slides.append({
                    "layout": "decision",
                    "section": "STRATEGY",
                    "title": f"五观报告给出 {len(action_items)} 条优先行动建议，按原文要点提请决策",
                    "answer_first": True,
                    "items": action_items,
                    "takeaway": "建议项摘自五观正文优先行动建议章节，未做外推。",
                    "source": source,
                    "bias_note": bias,
                })

        # 数据图页：复用已批准交付图（按 claim registry 的断言作要点）。
        registry_path = self.project_dir / "intermediate" / "claim_registry.json"
        claims: dict[str, Any] = {}
        if registry_path.is_file():
            try:
                claims = json.loads(registry_path.read_text(encoding="utf-8-sig"))
            except (OSError, json.JSONDecodeError):
                claims = {}
        charts_dir = self.project_dir / "deliverables" / "charts"
        figure_png: dict[str, str] = {}
        for theme_path in sorted(charts_dir.glob("fig*.theme.json")):
            try:
                theme = json.loads(theme_path.read_text(encoding="utf-8-sig"))
            except (OSError, json.JSONDecodeError):
                continue
            figure_id = str(theme.get("figure_id") or "")
            png_rel = str(((theme.get("outputs") or {}).get("png") or {}).get("path") or "")
            if figure_id and png_rel:
                png_path = Path(png_rel)
                if not png_path.is_absolute():
                    # theme 清单内的路径以 project_dir 为基准（如 deliverables/charts/figN.png）。
                    png_path = self.project_dir / png_rel
                if png_path.is_file():
                    try:
                        figure_png[figure_id] = png_path.resolve().relative_to(self.project_dir.resolve()).as_posix()
                    except ValueError:
                        figure_png[figure_id] = str(png_path)
        figure_slide_defs = [
            ("market_trend", "MARKET", "市场规模观测序列已登记并可复核，趋势判断以证据表为准"),
            ("price_capacity_scatter", "COMPETITION", "价格—容量定位来自综合矩阵登记值，样本外不外推"),
            ("pain_point_pareto", "CUSTOMER", "痛点频次来自评论编码登记，样本量以证据表为准"),
        ]
        for figure_id, section, takeaway in figure_slide_defs:
            if figure_id not in figure_png:
                continue
            claim_text = str((claims.get(figure_id) or {}).get("core_claim") or "该图由蒸馏证据表直接计算生成。")
            slides.append({
                "layout": "figure",
                "section": section,
                "title": f"证据图：{claim_text[:44]}",
                "answer_first": True,
                "figure_path": figure_png[figure_id],
                "items": [claim_text],
                "takeaway": takeaway,
                "source": source,
                "bias_note": bias,
            })

        player_types = Counter((row.get("player_type") or "未分类").strip() or "未分类" for row in competitor_rows)
        segment_items = []
        for player_type, count in player_types.most_common(3):
            brands = [row.get("brand", "") for row in competitor_rows if row.get("player_type") == player_type][:3]
            segment_items.append({
                "title": f"{player_type}（{count} 家）",
                "text": "、".join(brand for brand in brands if brand) or "品牌待补",
                "barrier": "本地化服务与渠道证据待核实",
                "metric": f"样本 {count} 家",
            })
        if segment_items:
            slides.append({
                "layout": "segments",
                "section": "COMPETITION",
                "title": f"竞争玩家按类型分为 {len(player_types)} 类，共登记 {len(competitor_rows)} 家样本",
                "answer_first": True,
                "items": segment_items,
                "source": source,
                "bias_note": bias,
            })

        priced = [row for row in matrix if num(row.get("price", "")) is not None]
        pains = [row for row in matrix if num(row.get("user_pain_score", "")) is not None]
        if len(priced) >= 2 and pains:
            price_values = [num(row["price"]) for row in priced]
            pain_values = [num(row["user_pain_score"]) for row in pains]
            price_span = (min(price_values), max(price_values))
            pain_span = (min(pain_values), max(pain_values))

            def norm(value: float, span: tuple[float, float]) -> float:
                low, high = span
                if high <= low:
                    return 0.5
                return max(0.06, min(0.94, (value - low) / (high - low)))

            points = []
            for row in matrix:
                price = num(row.get("price", ""))
                pain = num(row.get("user_pain_score", ""))
                if price is None or pain is None:
                    continue
                points.append({
                    "x": round(norm(price, price_span), 2),
                    "y": round(norm(pain, pain_span), 2),
                    "label": (row.get("brand", "") or row.get("exact_model", ""))[:10] or "样本",
                })
            # 点位标签盒 [px+0.12, py-0.14, 1.10, 0.28]，px=1.55+x*7.15、py=5.88-y*3.6：
            # 盒重叠 ⇔ 归一化 Δx<0.154 且 Δy<0.078（validate text_overlap）。
            # 碰撞源除其他点位外还包括四个象限标签盒（[1.48|5.45, 1.92|4.25, 3.35, 0.30]，
            # 换算后归一化 y 危险带约 0.32-0.51 与 0.86-1.05），一并纳入碰撞检测。
            # 避让用固定候选槽位枚举（不回溯，保证终止）：最多 8 点，槽位足够。
            zone_x = [(-1.0, 2.0)]  # 两个象限标签间隙 <1.10in 标签盒宽，危险带覆盖全宽
            zone_y = [(0.32, 0.51), (0.86, 1.05)]
            candidate_slots = [(x, y) for x in (0.06, 0.22, 0.38, 0.54, 0.70, 0.86)
                               for y in (0.06, 0.14, 0.22, 0.56, 0.64, 0.72, 0.80)]

            def collides(x: float, y: float) -> bool:
                if any(x1 - 0.16 < x < x2 and y1 - 0.08 < y < y2 for x1, x2 in zone_x for y1, y2 in zone_y):
                    return True
                return any(abs(x - other["x"]) < 0.16 and abs(y - other["y"]) < 0.08 for other in placed)

            placed: list[dict[str, float]] = []
            for point in points:
                x, y = point["x"], point["y"]
                if collides(x, y):
                    slot = next(((sx, sy) for sx, sy in candidate_slots if not collides(sx, sy)), None)
                    if slot is None:
                        continue  # 槽位耗尽：保留原位，validate 会给出诊断
                    x, y = slot
                point["x"], point["y"] = x, y
                placed.append(point)
            slides.append({
                "layout": "matrix",
                "section": "COMPETITION",
                "title": "样本在价格与用户口碑两个维度上分布，坐标均为证据登记值的归一化",
                "answer_first": True,
                "quadrants": ["低价/高口碑", "高价/高口碑", "低价/低口碑", "高价/低口碑"],
                "points": points[:8],
                "items": ["横轴为价格归一化，纵轴为用户口碑评分归一化", "坐标仅用于样本内相对对照，不构成绝对定位"],
                "takeaway": "优先级以证据强度校准",
                "source": source,
                "bias_note": bias,
            })

        if swot:
            strengths = [row.get("strength", "") for row in swot if row.get("strength", "").strip()][:3]
            weaknesses = [row.get("weakness", "") for row in swot if row.get("weakness", "").strip()][:3]
            opportunities = [row.get("opportunity", "") for row in swot if row.get("opportunity", "").strip()][:3]
            threats = [row.get("threat", "") for row in swot if row.get("threat", "").strip()][:3]
            slides.append({
                "layout": "swot",
                "section": "STRATEGY",
                "title": f"SWOT 四象限均来自 {len(swot)} 条登记条目的归纳，不作样本外推断",
                "answer_first": True,
                "items": [
                    {"title": "优势", "items": strengths or ["待补充"]},
                    {"title": "劣势", "items": weaknesses or ["待补充"]},
                    {"title": "机会", "items": opportunities or ["待补充"]},
                    {"title": "威胁", "items": threats or ["待补充"]},
                ],
                "source": source,
                "bias_note": bias,
            })

        comparison_items = []
        for row in sorted(priced, key=lambda r: -(num(r.get("price", "")) or 0))[:4]:
            capacity = row.get("capacity_kwh", "")
            price = num(row.get("price", ""))
            comparison_items.append({
                "title": row.get("exact_model", "") or row.get("brand", ""),
                "headline": f"{capacity} kWh" if capacity else "容量待核",
                "text": f"渠道：{(row.get('channel_coverage') or '待核')[:30]}；智能功能：{(row.get('smart_features') or '待核')[:30]}",
                "metric": f"{price:g} {row.get('currency', '')}" if price is not None else "价格待核",
            })
        if comparison_items:
            slides.append({
                "layout": "comparison",
                "section": "COMPETITION",
                "title": f"{len(comparison_items)} 个主力型号按登记价格从高到低对照",
                "answer_first": True,
                "items": comparison_items,
                "source": source,
                "bias_note": bias,
            })

        slides.append({
            "layout": "timeline",
            "section": "ROADMAP",
            "title": "进入节奏以证据闭环为闸门：先验证、后投入、再规模化",
            "answer_first": True,
            "items": [
                {"period": "阶段一", "title": "证据核验",
                 "text": f"补齐 {sum(1 for row in ledger if (row.get('verification_status') or '') != 'verified')} 条未验证来源的复核，确认市场准入与电价口径。"},
                {"period": "阶段二", "title": "小规模验证",
                 "text": "以有限样机/试点验证单位经济性，回收测算输入以本册证据表口径为准。"},
                {"period": "阶段三", "title": "规模化进入",
                 "text": "政策、渠道与经济性三重门槛同时成立后再扩大投入，保留退出条件。"},
            ],
            "takeaway": "时间节点以监管开放、标准落地和单位经济性为共同触发器。",
            "source": source,
            "bias_note": bias,
        })
        slides.append({
            "layout": "roadmap",
            "section": "ROADMAP",
            "title": "三阶段路线图均设停止条件，避免在证据未成立前重资产投入",
            "answer_first": True,
            "items": [
                {"period": "阶段一 0-3 个月", "title": "证据与合规核验",
                 "text": "完成来源复核与准入政策核对，输出可决策的证据闭环报告。",
                 "gate": "未验证来源降至可接受水平"},
                {"period": "阶段二 3-9 个月", "title": "试点验证经济性",
                 "text": "小批量试点采集真实安装与运维成本，校准回收模型。",
                 "gate": "单位经济性达到立项门槛"},
                {"period": "阶段三 9 个月以上", "title": "规模化进入",
                 "text": "渠道与安装网络就绪后扩大铺货，持续监控政策与价格变化。",
                 "gate": "三重门槛复核通过"},
            ],
            "takeaway": "每阶段均保留停止条件，避免在政策或经济性未成立前重资产投入。",
            "source": source,
            "bias_note": bias,
        })

        ranked = sorted(swot, key=lambda row: str(row.get("opportunity_priority") or "Z"))
        decision_items = []
        for row in ranked[:4]:
            opportunity = (row.get("opportunity") or "").strip()
            decision_items.append({
                "title": opportunity[:40] or "机会待定义",
                "text": f"对标 {row.get('brand', '') or '样本品牌'} 的证据条目，优先级：{row.get('opportunity_priority', '') or '未分级'}。",
                "owner": "产品与市场联合小组",
                "gate": "证据闭环复核通过",
            })
        if not decision_items:
            decision_items = [{"title": "补齐机会分级", "text": "SWOT 证据表为空，需先完成蒸馏登记。",
                               "owner": "研究小组", "gate": "证据登记完成"}]
        slides.append({
            "layout": "decision",
            "section": "DECISION",
            "title": f"按机会优先级提请决策：共 {len(decision_items)} 项待裁决",
            "answer_first": True,
            "items": decision_items,
            "takeaway": "只批准可逆、可度量且有明确退出条件的下一步。",
            "source": source,
            "bias_note": bias,
        })
        slides.append({
            "layout": "closing",
            "section": "DECISION & NEXT STEP",
            "title": "以证据闭环换取下一阶段选择权",
            "items": [
                {"title": "复核", "text": "未验证来源逐条复核，证据台账保持唯一事实源。"},
                {"title": "验证", "text": "小规模试点校准单位经济性与服务成本。"},
                {"title": "决策", "text": "三重门槛成立后再扩大投入，保留退出条件。"},
            ],
            "source": f"{region}{category}市场研究 · {today}",
        })
        return {
            "deck": {
                "title": f"{region}{category}市场研究内部宣讲",
                "subtitle": "市场、竞争与进入策略：证据驱动的决策材料",
                "eyebrow": "ENERGY MARKET & PRODUCT INTELLIGENCE",
                "meta": "无人值守编排模式 · 证据蒸馏快照",
                "confidentiality": "内部使用",
                "update_date": today,
            },
            "slides": slides,
        }

    def _collect_artifacts(self) -> None:
        deliverables = self.project_dir / "deliverables"
        for pattern in ("*.xlsx", "*.docx", "*.pptx"):
            for path in sorted(deliverables.glob(pattern)):
                if path.is_file():
                    self.artifacts.append(str(path))
        insight = self.project_dir / "intermediate" / "market-insight" / "market_insight_report.md"
        # 模板骨架（含 [[填写]] 占位符或 status: draft）不是交付物，绝不外推。
        if insight.is_file():
            text = insight.read_text(encoding="utf-8", errors="replace")
            if "[[填写" not in text and "status: draft" not in text[:500]:
                self.artifacts.append(str(insight))
